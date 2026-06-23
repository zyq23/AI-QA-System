from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document

from app.config import Settings
from app.domain import DraftAnswer, QueryAnalysis, RetrievalHit, ReviewResult
from app.services.llm import LlmService
from app.services.ml import EmbeddingService, RerankerService
from app.services.ragflow import RagflowClient
from app.services.evaluation import EvaluationService
from app.services.retrieval import AdaptiveRetrievalService, FallbackRetrievalService, RagflowRetrievalService, RetrievalResult, RetrievalService


def build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("宇树 G1 知识卡", level=1)
    document.add_paragraph("宇树 G1 支持资料问答原型验证。")
    document.add_paragraph("如果资料不足，系统必须提示当前知识库没有直接证据。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_hit(text: str, *, section_path: str = "测试") -> RetrievalHit:
    return RetrievalHit(
        chunk_id="chunk-1",
        document_id="doc-1",
        version_id="ver-1",
        file_name="kb.docx",
        page_or_slide="docx",
        section_path=section_path,
        snippet=text,
        markdown_text=text,
        plain_text=text,
        trust_level="internal",
        source_type="upload",
        fusion_score=1.0,
        rerank_score=1.0,
    )


def test_upload_and_chat_flow(client):
    files = [("files", ("g1.docx", build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    response = client.post(
        "/api/admin/documents",
        headers={"X-Admin-Token": "test-token"},
        files=files,
        data={"source_type": "upload", "trust_level": "internal"},
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["document_ids"]

    docs = client.get("/api/admin/documents", headers={"X-Admin-Token": "test-token"})
    assert docs.status_code == 200
    listed = docs.json()
    assert listed[0]["chunk_count"] >= 1

    chat = client.post(
        "/api/chat/query",
        json={"question": "系统在资料不足时应该怎么回答？"},
    )
    assert chat.status_code == 200, chat.text
    answer = chat.json()
    assert answer["citations"]
    assert "知识库没有直接证据" in answer["answer"] or "知识库没有直接证据" in answer["grounded_answer"]
    assert "文件：" not in answer["answer"]
    assert "章节：" not in answer["answer"]
    assert "g1.docx" not in answer["answer"]


def test_admin_can_run_eval_job(client, app_env: Path):
    eval_dataset = app_env / "data" / "evals" / "cases.json"
    eval_dataset.parent.mkdir(parents=True, exist_ok=True)
    eval_dataset.write_text(
        json.dumps(
            [
                {
                    "id": "oos_only",
                    "category": "out_of_scope",
                    "question": "今天上海天气怎么样？",
                    "expected_insufficient": True,
                    "expected_grounded": False,
                }
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    response = client.post("/api/admin/evals/run", headers={"X-Admin-Token": "test-token"})
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["status"] == "queued"

    jobs = client.app.state.container.repository.list_jobs()
    eval_job = next(job for job in jobs if job["id"] == payload["job_id"])
    assert eval_job["job_type"] == "evaluation"
    assert eval_job["status"] == "completed"
    assert eval_job["result"]["summary"]["pass_rate"] == 1.0
    assert "direct_answer_rate" in eval_job["result"]["summary"]
    assert "no_leak_rate" in eval_job["result"]["summary"]
    assert "concise_rate" in eval_job["result"]["summary"]
    assert "reviewer_intervention_rate" in eval_job["result"]["summary"]
    assert "quality_issue_breakdown" in eval_job["result"]["summary"]
    assert "failed_check_breakdown" in eval_job["result"]["summary"]
    assert Path(eval_job["result"]["report_path"]).exists()


def test_admin_can_run_eval_job_with_dataset_override(client, app_env: Path):
    client.app.state.container.settings.eval_api_base_url = None
    quick_dataset = app_env / "data" / "evals" / "quick_cases.json"
    quick_dataset.parent.mkdir(parents=True, exist_ok=True)
    quick_dataset.write_text(
        json.dumps(
            [
                {
                    "id": "quick_case",
                    "category": "out_of_scope",
                    "question": "库外问题怎么处理？",
                    "expected_insufficient": True,
                    "expected_grounded": False,
                }
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    response = client.post(
        "/api/admin/evals/run?dataset=data/evals/quick_cases.json",
        headers={"X-Admin-Token": "test-token"},
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["status"] == "queued"

    jobs = client.app.state.container.repository.list_jobs()
    eval_job = next(job for job in jobs if job["id"] == payload["job_id"])
    assert eval_job["payload"]["dataset"].endswith("data/evals/quick_cases.json")
    assert eval_job["result"]["summary"]["total_turns"] == 1


def test_admin_jobs_api_returns_job_summary(client):
    eval_dataset = Path(client.app.state.container.settings.eval_dataset_path)
    eval_dataset.parent.mkdir(parents=True, exist_ok=True)
    eval_dataset.write_text(
        json.dumps(
            [
                {
                    "id": "jobs_api_smoke",
                    "category": "out_of_scope",
                    "question": "库外问题怎么处理？",
                    "expected_insufficient": True,
                    "expected_grounded": False,
                }
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    response = client.post("/api/admin/evals/run", headers={"X-Admin-Token": "test-token"})
    assert response.status_code == 200, response.text

    jobs = client.get("/api/admin/jobs", headers={"X-Admin-Token": "test-token"})
    assert jobs.status_code == 200, jobs.text
    payload = jobs.json()
    assert payload
    assert "job_id" in payload[0]
    assert "id" not in payload[0]


def test_repository_latest_job_skips_running_evaluation_with_newer_non_eval_jobs(app_env: Path):
    from app.db import Database
    from app.repositories import Repository

    db_path = app_env / "data" / "runtime" / "app.db"
    db_path.parent.mkdir(parents=True, exist_ok=True)
    db = Database(db_path)
    db.initialize()
    repo = Repository(db)

    running = repo.create_job("evaluation", {"dataset": "data/evals/knowledge_base_eval_cases.json"})
    repo.update_job(running["id"], status="running", message="still running")

    completed = repo.create_job("evaluation", {"dataset": "data/evals/ppt_company_p0_8.json"})
    repo.update_job(completed["id"], status="completed", result={"summary": {"pass_rate": 1.0}})

    ingest = repo.create_job("ingest_document", {"filename": "demo.pptx"})
    repo.update_job(ingest["id"], status="completed")

    latest_eval = repo.latest_job(job_type="evaluation", status="completed")

    assert latest_eval is not None
    assert latest_eval["id"] == completed["id"]
    assert latest_eval["status"] == "completed"
    assert latest_eval["payload"]["dataset"] == "data/evals/ppt_company_p0_8.json"


def test_answer_hides_document_title_even_for_source_question(client):
    files = [("files", ("g1.docx", build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    response = client.post(
        "/api/admin/documents",
        headers={"X-Admin-Token": "test-token"},
        files=files,
        data={"source_type": "upload", "trust_level": "internal"},
    )
    assert response.status_code == 200, response.text

    chat = client.post("/api/chat/query", json={"question": "哪份资料提到了知识卡内容？"})
    assert chat.status_code == 200, chat.text
    answer = chat.json()["answer"]
    assert "g1.docx" not in answer
    assert "宇树 G1 知识卡" not in answer
    assert "相关资料这份资料" not in answer


def test_answer_runs_api_exposes_quality_trace(client):
    files = [("files", ("g1.docx", build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    response = client.post(
        "/api/admin/documents",
        headers={"X-Admin-Token": "test-token"},
        files=files,
        data={"source_type": "upload", "trust_level": "internal"},
    )
    assert response.status_code == 200, response.text

    chat = client.post("/api/chat/query", json={"question": "哪份资料提到了知识卡内容？"})
    assert chat.status_code == 200, chat.text
    answer_run_id = chat.json()["answer_run_id"]
    assert answer_run_id

    listing = client.get("/api/admin/answer-runs", headers={"X-Admin-Token": "test-token"})
    assert listing.status_code == 200, listing.text
    rows = listing.json()
    assert rows
    assert rows[0]["answer_run_id"] == answer_run_id
    assert rows[0]["question_type"] == "factoid"
    assert "expanded_query" in rows[0]
    assert "expansion_terms" in rows[0]
    assert "content_policy_blocked" in rows[0]

    detail = client.get(f"/api/admin/answer-runs/{answer_run_id}", headers={"X-Admin-Token": "test-token"})
    assert detail.status_code == 200, detail.text
    payload = detail.json()
    assert payload["draft"]
    assert payload["review"]
    assert payload["retrieval"]
    assert payload["final_answer"]
    assert "expanded_query" in payload
    assert "expansion_terms" in payload
    assert "content_policy_blocked" in payload
    assert "g1.docx" not in payload["final_answer"]


def test_robot_query_returns_robot_safe_payload(client):
    files = [("files", ("g1.docx", build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    response = client.post(
        "/api/admin/documents",
        headers={"X-Admin-Token": "test-token"},
        files=files,
        data={"source_type": "upload", "trust_level": "internal"},
    )
    assert response.status_code == 200, response.text

    robot = client.post(
        "/api/robot/query",
        json={"question": "系统在资料不足时应该怎么回答？", "client_id": "g1-edu-dock"},
    )
    assert robot.status_code == 200, robot.text
    payload = robot.json()
    assert payload["answer"]
    assert payload["tts_text"]
    assert payload["conversation_id"]
    assert "citations" not in payload
    assert "grounded_answer" not in payload
    assert "inference_note" not in payload


def test_robot_query_reuses_conversation_id(client):
    files = [("files", ("g1.docx", build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    response = client.post(
        "/api/admin/documents",
        headers={"X-Admin-Token": "test-token"},
        files=files,
        data={"source_type": "upload", "trust_level": "internal"},
    )
    assert response.status_code == 200, response.text

    first = client.post("/api/robot/query", json={"question": "系统在资料不足时应该怎么回答？"})
    assert first.status_code == 200, first.text
    conversation_id = first.json()["conversation_id"]

    second = client.post(
        "/api/robot/query",
        json={"question": "那它会不会直接编造答案？", "conversation_id": conversation_id},
    )
    assert second.status_code == 200, second.text
    assert second.json()["conversation_id"] == conversation_id


def test_admin_retrieval_test_api_returns_hits(client):
    files = [("files", ("g1.docx", build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    response = client.post(
        "/api/admin/documents",
        headers={"X-Admin-Token": "test-token"},
        files=files,
        data={"source_type": "upload", "trust_level": "internal"},
    )
    assert response.status_code == 200, response.text

    debug = client.post(
        "/api/admin/retrieval/test",
        headers={"X-Admin-Token": "test-token"},
        json={"question": "资料不足时系统应该怎么回答？", "top_k": 5},
    )
    assert debug.status_code == 200, debug.text
    payload = debug.json()
    assert payload["question"] == "资料不足时系统应该怎么回答？"
    assert "expanded_query" in payload
    assert payload["hits"]
    assert "grounded" in payload
    assert payload["backend_path"] == "local"
    assert "route_reason" in payload
    assert "remote_attempted" in payload
    assert "local_top_score" in payload
    assert "local_quality_score" in payload
    assert "remote_quality_score" in payload
    assert "local_grounded_score_threshold" in payload
    assert "score" in payload["hits"][0]
    assert "section_path" in payload["hits"][0]
    assert "keyword_rank" in payload["hits"][0]
    assert "vector_rank" in payload["hits"][0]
    assert "fusion_score" in payload["hits"][0]
    assert "rerank_score" in payload["hits"][0]
    assert "focus_matches" in payload["hits"][0]


def test_ragflow_client_lists_documents_and_chunks():
    def handler(request: httpx.Request) -> httpx.Response:
        if request.url.path == "/api/v1/datasets/ds-1/documents":
            return httpx.Response(200, json={"code": 0, "data": {"docs": [{"id": "doc-1", "name": "知识库材料.pdf"}]}})
        if request.url.path == "/api/v1/datasets/ds-1/documents/doc-1/chunks":
            return httpx.Response(
                200,
                json={"code": 0, "data": {"chunks": [{"id": "c1", "content": "第一段知识点", "positions": ["page 1"]}]}},
            )
        return httpx.Response(404, json={"code": 404, "message": "not found"})

    client = RagflowClient(
        base_url="http://ragflow.local",
        api_key="token",
        client=httpx.Client(base_url="http://ragflow.local", transport=httpx.MockTransport(handler)),
    )
    docs = client.list_dataset_documents("ds-1")
    chunks = client.list_document_chunks("ds-1", "doc-1")

    assert docs[0]["name"] == "知识库材料.pdf"
    assert chunks[0]["content"] == "第一段知识点"


def test_query_expansion_adds_domain_synonyms():
    expanded_query, expansion_terms, expanded_focus_terms = RetrievalService._expand_query(
        "产业学院的认证覆盖哪些级别？",
        ["产业学院的认证覆盖", "级别"],
    )
    assert "HCIA" in expansion_terms
    assert "HCIP" in expansion_terms
    assert "HCIE" in expansion_terms
    assert "等级" in expansion_terms
    assert "根技术认证运营" in expansion_terms
    assert "HCIA" in expanded_query
    assert "HCIA" in expanded_focus_terms

    expanded_query2, expansion_terms2, _ = RetrievalService._expand_query("产业学院采用哪种管理模式？", None)
    assert "治理模式" in expansion_terms2
    assert "治理模式" in expanded_query2

    expanded_query3, expansion_terms3, _ = RetrievalService._expand_query("产业学院有哪些沟通制度？", None)
    assert "月例会" in expansion_terms3
    assert "季汇报" in expansion_terms3
    assert "年总结" in expansion_terms3
    assert "沟通机制" in expanded_query3

    expanded_query4, expansion_terms4, _ = RetrievalService._expand_query("Edge智控APP绑定失败时应该先检查什么？", None)
    assert "互联网" in expansion_terms4
    assert "网络认证" in expansion_terms4
    assert "无法绑定设备" in expanded_query4

    expanded_query5, expansion_terms5, _ = RetrievalService._expand_query("协作式机械臂产品的开放性实验环境是什么？", None)
    assert "Jupyter Notebook" in expansion_terms5
    assert "Jupyter Notebook" in expanded_query5

    expanded_query6, expansion_terms6, _ = RetrievalService._expand_query("产业学院的最高决策单位是什么？", None)
    assert "最高决策机构" in expansion_terms6
    assert "决策层" in expanded_query6

    expanded_query7, expansion_terms7, _ = RetrievalService._expand_query("产业学院有哪些协同机制？", None)
    assert "沟通制度" in expansion_terms7
    assert "沟通机制" in expanded_query7

    expanded_query8, expansion_terms8, _ = RetrievalService._expand_query("华为ICT学院课程体系有哪些？", None)
    assert "课程资源" in expansion_terms8
    assert "课程类型" in expanded_query8

    expanded_query9, expansion_terms9, _ = RetrievalService._expand_query("根技术课程体系", None)
    assert "17个学院70个专业" in expansion_terms9
    assert "AIGC实战平台" in expanded_query9

    expanded_query10, expansion_terms10, _ = RetrievalService._expand_query("协作式机械臂产品适用哪些课程？", None)
    assert "适用课程" in expansion_terms10
    assert "Python程序设计" in expansion_terms10
    assert "深度学习" in expanded_query10

    expanded_query11, expansion_terms11, _ = RetrievalService._expand_query("宇树G1机器人的关节数量是多少？", None)
    assert "Unitree G1" in expansion_terms11
    assert "总自由度" in expanded_query11

    expanded_query12, expansion_terms12, _ = RetrievalService._expand_query("华为根技术体验中心的核心定位主线是什么？", None)
    assert "根技术筑基" in expansion_terms12
    assert "产教融育人" in expansion_terms12
    assert "师范践初心" in expanded_query12

    expanded_query13, expansion_terms13, _ = RetrievalService._expand_query("产业学院提到的AI核心课程是什么？", None)
    assert "现代教育技术与智慧教学" in expansion_terms13
    assert "AI 赋能核心课程协同共建" in expanded_query13

    expanded_query14, expansion_terms14, _ = RetrievalService._expand_query("协作机器人的额定负载是多少？", None)
    assert "3kg" in expansion_terms14
    assert "主要硬件参数" in expanded_query14

    expanded_query15, expansion_terms15, _ = RetrievalService._expand_query("华为在根技术研发布局是什么？", None)
    assert "强力投入研究与开发" in expansion_terms15
    assert "标准与专利" in expanded_query15

    expanded_query16, expansion_terms16, _ = RetrievalService._expand_query("华为通过哪三个重构，哪五大方向突围？", None)
    assert "理论重构" in expansion_terms16
    assert "运营系统" in expanded_query16

    expanded_query17, expansion_terms17, _ = RetrievalService._expand_query("怎么样提交申请？", None)
    assert "申请步骤" in expansion_terms17
    assert "华为审核" in expanded_query17


def test_settings_read_ragflow_dataset_ids_from_env_aliases():
    settings = Settings.model_validate(
        {
            "RAGFLOW_DATASET_IDS": "dataset-1,dataset-2",
            "RAGFLOW_DOCUMENT_IDS": "doc-1,doc-2",
        }
    )
    assert settings.ragflow_dataset_ids == ["dataset-1", "dataset-2"]
    assert settings.ragflow_document_ids == ["doc-1", "doc-2"]


def test_embedding_service_caches_repeated_query_vectors():
    service = EmbeddingService(model_name="stub", use_stub=True, query_cache_size=8)

    call_count = {"count": 0}

    def fake_embed_documents(texts):
        call_count["count"] += 1
        return [[float(len(texts[0]))]]

    service.embed_documents = fake_embed_documents  # type: ignore[method-assign]

    first = service.embed_query("重复查询")
    second = service.embed_query("重复查询")

    assert first == second
    assert call_count["count"] == 1


def test_reranker_service_caches_repeated_scores():
    service = RerankerService(model_name="stub", use_stub=False, score_cache_size=8)

    class FakeBackend:
        def __init__(self) -> None:
            self.calls = 0

        def compute_score(self, pairs, batch_size=8):
            self.calls += 1
            return [float(len(text)) for _, text in pairs]

    backend = FakeBackend()
    service._backend = backend

    hits = [build_hit("第一段证据"), build_hit("第二段证据", section_path="第二段")]
    first = service.rerank("重复问题", hits)
    second = service.rerank("重复问题", hits)

    assert backend.calls == 1
    assert [hit.rerank_score for hit in first] == [hit.rerank_score for hit in second]


def test_retrieval_marks_noisy_ocr_blocks():
    assert RetrievalService._looks_like_noisy_ocr("金1 银合金！！ 错会金1 ??? Unitree G1")
    assert not RetrievalService._looks_like_noisy_ocr("适用课程包括Python程序设计、深度学习和机器视觉。")


def test_grounded_allows_split_multichunk_factoids():
    hits = [
        build_hit("华为通过3个重构，5大方向突围，构建根技术，实现科技自立自强。", section_path="slide-5 / body"),
        build_hit("理论重构 架构重构 软件重构", section_path="slide-5 / list"),
        build_hit("基础理论 基础硬件 基础软件 开发工具 运营系统", section_path="slide-5 / list-2"),
    ]
    hits[0].rerank_score = 4.2
    hits[1].rerank_score = 0.8
    hits[2].rerank_score = 0.8

    assert RetrievalService._grounded("华为通过哪三个重构，哪五大方向突围", hits, ["理论重构", "架构重构"])


def test_retrieval_groups_sibling_ppt_chunks_for_foundation_model_questions():
    class FakeRepository:
        def __init__(self, rows):
            self.rows = rows

        def keyword_search(self, query: str, limit: int):
            return []

        def get_chunks_by_ids(self, chunk_ids):
            mapping = {row["id"]: row for row in self.rows}
            return [mapping[chunk_id] for chunk_id in chunk_ids if chunk_id in mapping]

    class FakeEmbeddingService:
        def embed_query(self, text: str):
            return [0.1, 0.2]

    class FakeVectorStore:
        def query(self, query_embedding, n_results: int):
            return [
                {"chunk_id": "new-slide19-ocr", "distance": 0.01},
                {"chunk_id": "old-slide8-model", "distance": 0.02},
                {"chunk_id": "new-slide19-model", "distance": 0.03},
                {"chunk_id": "new-slide19-body", "distance": 0.04},
            ]

    class FakeReranker:
        def rerank(self, query: str, hits: list[RetrievalHit]) -> list[RetrievalHit]:
            base_scores = {
                "（OCR、语音识别等）": 0.8,
                "大模型基础应用": 1.0,
                "通用大模型": 0.2,
                "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据": 0.1,
            }
            reranked = []
            for hit in hits:
                reranked.append(
                    RetrievalHit(
                        chunk_id=hit.chunk_id,
                        document_id=hit.document_id,
                        version_id=hit.version_id,
                        file_name=hit.file_name,
                        page_or_slide=hit.page_or_slide,
                        section_path=hit.section_path,
                        snippet=hit.snippet,
                        markdown_text=hit.markdown_text,
                        plain_text=hit.plain_text,
                        trust_level=hit.trust_level,
                        source_type=hit.source_type,
                        fusion_score=hit.fusion_score,
                        rerank_score=base_scores[hit.plain_text],
                        raw_scores=hit.raw_scores,
                    )
                )
            return reranked

    rows = [
        {
            "id": "new-slide19-ocr",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body-2",
            "plain_text": "（OCR、语音识别等）",
            "markdown_text": "（OCR、语音识别等）",
        },
        {
            "id": "new-slide19-model",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body-1",
            "plain_text": "通用大模型",
            "markdown_text": "通用大模型",
        },
        {
            "id": "new-slide19-body",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body",
            "plain_text": "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
            "markdown_text": "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
        },
        {
            "id": "old-slide8-model",
            "document_id": "doc-old",
            "version_id": "ver-old",
            "file_name": "根技术体验中心展厅内涵建设v2-cx.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-8",
            "section_path": "幻灯片 8 / 幻灯片 8 / picture-ocr-69",
            "plain_text": "大模型基础应用",
            "markdown_text": "大模型基础应用",
        },
    ]

    service = RetrievalService(
        FakeRepository(rows),
        FakeEmbeddingService(),
        FakeReranker(),
        FakeVectorStore(),
        candidates=8,
        default_top_k=4,
        retrieval_mode="hybrid",
    )

    result = service.retrieve("基础模型页除了通用大模型，还列了哪些感知或解析能力？", top_k=4)

    assert [hit.page_or_slide for hit in result.hits[:3]] == ["slide-19", "slide-19", "slide-19"]
    top3_texts = [hit.plain_text for hit in result.hits[:3]]
    assert "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据" in top3_texts
    assert "（OCR、语音识别等）" in top3_texts
    assert "通用大模型" in top3_texts
    assert result.hits[3].plain_text == "大模型基础应用"


def test_retrieval_promotes_foundation_model_family_and_capability_chunks():
    class FakeRepository:
        def __init__(self, rows):
            self.rows = rows

        def keyword_search(self, query: str, limit: int):
            return []

        def get_chunks_by_ids(self, chunk_ids):
            mapping = {row["id"]: row for row in self.rows}
            return [mapping[chunk_id] for chunk_id in chunk_ids if chunk_id in mapping]

    class FakeEmbeddingService:
        def embed_query(self, text: str):
            return [0.1, 0.2]

    class FakeVectorStore:
        def query(self, query_embedding, n_results: int):
            return [
                {"chunk_id": "new-slide19-body", "distance": 0.01},
                {"chunk_id": "old-voice", "distance": 0.02},
                {"chunk_id": "new-slide19-ocr", "distance": 0.03},
                {"chunk_id": "new-slide19-model", "distance": 0.04},
                {"chunk_id": "new-slide19-family", "distance": 0.05},
            ]

    class FakeReranker:
        def rerank(self, query: str, hits: list[RetrievalHit]) -> list[RetrievalHit]:
            base_scores = {
                "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据": 0.6,
                "利用开放平台进行语音识别": 1.0,
                "（OCR、语音识别等）": 0.4,
                "通用大模型": 0.2,
                "（deepseek、通义千问、文心一言等）": 0.1,
            }
            reranked = []
            for hit in hits:
                reranked.append(
                    RetrievalHit(
                        chunk_id=hit.chunk_id,
                        document_id=hit.document_id,
                        version_id=hit.version_id,
                        file_name=hit.file_name,
                        page_or_slide=hit.page_or_slide,
                        section_path=hit.section_path,
                        snippet=hit.snippet,
                        markdown_text=hit.markdown_text,
                        plain_text=hit.plain_text,
                        trust_level=hit.trust_level,
                        source_type=hit.source_type,
                        fusion_score=hit.fusion_score,
                        rerank_score=base_scores[hit.plain_text],
                        raw_scores=hit.raw_scores,
                    )
                )
            return reranked

    rows = [
        {
            "id": "new-slide19-body",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body",
            "plain_text": "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
            "markdown_text": "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
        },
        {
            "id": "old-voice",
            "document_id": "doc-old",
            "version_id": "ver-old",
            "file_name": "广东技术师范大学-根技术人才培养合作汇报-0.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-28",
            "section_path": "幻灯片 28 / 幻灯片 28 / body-22",
            "plain_text": "利用开放平台进行语音识别",
            "markdown_text": "利用开放平台进行语音识别",
        },
        {
            "id": "new-slide19-ocr",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body-2",
            "plain_text": "（OCR、语音识别等）",
            "markdown_text": "（OCR、语音识别等）",
        },
        {
            "id": "new-slide19-model",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body-1",
            "plain_text": "通用大模型",
            "markdown_text": "通用大模型",
        },
        {
            "id": "new-slide19-family",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body-2",
            "plain_text": "（deepseek、通义千问、文心一言等）",
            "markdown_text": "（deepseek、通义千问、文心一言等）",
        },
    ]

    service = RetrievalService(
        FakeRepository(rows),
        FakeEmbeddingService(),
        FakeReranker(),
        FakeVectorStore(),
        candidates=8,
        default_top_k=5,
        retrieval_mode="hybrid",
    )

    result = service.retrieve("基础模型页提到的平台能力包含哪些模型或数据治理能力？", top_k=5)

    top4_texts = [hit.plain_text for hit in result.hits[:4]]
    top3_texts = [hit.plain_text for hit in result.hits[:3]]
    assert "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据" in top4_texts
    assert "（OCR、语音识别等）" in top4_texts
    assert "（deepseek、通义千问、文心一言等）" in top4_texts
    assert "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据" in top3_texts
    assert "（OCR、语音识别等）" in top3_texts
    assert "（deepseek、通义千问、文心一言等）" in top3_texts
    assert "利用开放平台进行语音识别" not in top3_texts


def test_retrieval_demotes_model_family_chunks_for_capability_only_foundation_model_question():
    class FakeRepository:
        def __init__(self, rows):
            self.rows = rows

        def keyword_search(self, query: str, limit: int):
            return []

        def get_chunks_by_ids(self, chunk_ids):
            mapping = {row["id"]: row for row in self.rows}
            return [mapping[chunk_id] for chunk_id in chunk_ids if chunk_id in mapping]

    class FakeEmbeddingService:
        def embed_query(self, text: str):
            return [0.1, 0.2]

    class FakeVectorStore:
        def query(self, query_embedding, n_results: int):
            return [
                {"chunk_id": "new-slide19-family", "distance": 0.01},
                {"chunk_id": "old-model", "distance": 0.02},
                {"chunk_id": "new-slide19-model", "distance": 0.03},
                {"chunk_id": "new-slide19-body", "distance": 0.04},
                {"chunk_id": "new-slide19-ocr", "distance": 0.05},
            ]

    class FakeReranker:
        def rerank(self, query: str, hits: list[RetrievalHit]) -> list[RetrievalHit]:
            base_scores = {
                "（deepseek、通义千问、文心一言等）": 1.0,
                "大模型基础应用": 0.9,
                "通用大模型": 0.6,
                "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据": 0.4,
                "（OCR、语音识别等）": 0.3,
            }
            reranked = []
            for hit in hits:
                reranked.append(
                    RetrievalHit(
                        chunk_id=hit.chunk_id,
                        document_id=hit.document_id,
                        version_id=hit.version_id,
                        file_name=hit.file_name,
                        page_or_slide=hit.page_or_slide,
                        section_path=hit.section_path,
                        snippet=hit.snippet,
                        markdown_text=hit.markdown_text,
                        plain_text=hit.plain_text,
                        trust_level=hit.trust_level,
                        source_type=hit.source_type,
                        fusion_score=hit.fusion_score,
                        rerank_score=base_scores[hit.plain_text],
                        raw_scores=hit.raw_scores,
                    )
                )
            return reranked

    rows = [
        {
            "id": "new-slide19-family",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body-2",
            "plain_text": "（deepseek、通义千问、文心一言等）",
            "markdown_text": "（deepseek、通义千问、文心一言等）",
        },
        {
            "id": "old-model",
            "document_id": "doc-old",
            "version_id": "ver-old",
            "file_name": "根技术体验中心展厅内涵建设v2-cx.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-8",
            "section_path": "幻灯片 8 / 幻灯片 8 / picture-ocr-69",
            "plain_text": "大模型基础应用",
            "markdown_text": "大模型基础应用",
        },
        {
            "id": "new-slide19-model",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body-1",
            "plain_text": "通用大模型",
            "markdown_text": "通用大模型",
        },
        {
            "id": "new-slide19-body",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body",
            "plain_text": "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
            "markdown_text": "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
        },
        {
            "id": "new-slide19-ocr",
            "document_id": "doc-new",
            "version_id": "ver-new",
            "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
            "source_type": "bootstrap",
            "trust_level": "internal",
            "page_or_slide": "slide-19",
            "section_path": "幻灯片 19 / 幻灯片 19 / body-2",
            "plain_text": "（OCR、语音识别等）",
            "markdown_text": "（OCR、语音识别等）",
        },
    ]

    service = RetrievalService(
        FakeRepository(rows),
        FakeEmbeddingService(),
        FakeReranker(),
        FakeVectorStore(),
        candidates=8,
        default_top_k=5,
        retrieval_mode="hybrid",
    )

    result = service.retrieve("基础模型页除了通用大模型，还列了哪些感知或解析能力？", top_k=5)

    top3_texts = [hit.plain_text for hit in result.hits[:3]]
    assert "（OCR、语音识别等）" in top3_texts
    assert "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据" in top3_texts
    assert "（deepseek、通义千问、文心一言等）" not in top3_texts


def test_retrieval_widens_rerank_window_for_high_focus_foundation_model_enumeration():
    class FakeRepository:
        def __init__(self, rows):
            self.rows = rows

        def keyword_search(self, query: str, limit: int):
            return []

        def get_chunks_by_ids(self, chunk_ids):
            mapping = {row["id"]: row for row in self.rows}
            return [mapping[chunk_id] for chunk_id in chunk_ids if chunk_id in mapping]

    class FakeEmbeddingService:
        def embed_query(self, text: str):
            return [0.1, 0.2]

    class FakeVectorStore:
        def query(self, query_embedding, n_results: int):
            # The key capability body is intentionally placed beyond top_k * 3 (=18 for top_k=6).
            return [
                {"chunk_id": f"filler-{idx:02d}", "distance": 0.01 + idx * 0.001}
                for idx in range(1, 19)
            ] + [
                {"chunk_id": "slide19-body", "distance": 0.04},
                {"chunk_id": "slide19-ocr", "distance": 0.041},
                {"chunk_id": "slide19-model", "distance": 0.042},
            ]

    class RecordingReranker:
        def __init__(self):
            self.seen_texts: list[str] = []

        def rerank(self, query: str, hits: list[RetrievalHit]) -> list[RetrievalHit]:
            self.seen_texts = [hit.plain_text for hit in hits]
            base_scores = {
                "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据": 0.9,
                "（OCR、语音识别等）": 0.7,
                "通用大模型": 0.4,
            }
            reranked = []
            for hit in hits:
                reranked.append(
                    RetrievalHit(
                        chunk_id=hit.chunk_id,
                        document_id=hit.document_id,
                        version_id=hit.version_id,
                        file_name=hit.file_name,
                        page_or_slide=hit.page_or_slide,
                        section_path=hit.section_path,
                        snippet=hit.snippet,
                        markdown_text=hit.markdown_text,
                        plain_text=hit.plain_text,
                        trust_level=hit.trust_level,
                        source_type=hit.source_type,
                        fusion_score=hit.fusion_score,
                        rerank_score=base_scores.get(hit.plain_text, -1.0),
                        raw_scores=hit.raw_scores,
                    )
                )
            return reranked

    rows = []
    for idx in range(1, 19):
        rows.append(
            {
                "id": f"filler-{idx:02d}",
                "document_id": "doc-old",
                "version_id": "ver-old",
                "file_name": f"old-{idx:02d}.pptx",
                "source_type": "bootstrap",
                "trust_level": "internal",
                "page_or_slide": f"slide-{idx}",
                "section_path": f"幻灯片 {idx} / body",
                "plain_text": f"填充块 {idx}",
                "markdown_text": f"填充块 {idx}",
            }
        )
    rows.extend(
        [
            {
                "id": "slide19-body",
                "document_id": "doc-new",
                "version_id": "ver-new",
                "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
                "source_type": "bootstrap",
                "trust_level": "internal",
                "page_or_slide": "slide-19",
                "section_path": "幻灯片 19 / 幻灯片 19 / body",
                "plain_text": "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
                "markdown_text": "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
            },
            {
                "id": "slide19-ocr",
                "document_id": "doc-new",
                "version_id": "ver-new",
                "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
                "source_type": "bootstrap",
                "trust_level": "internal",
                "page_or_slide": "slide-19",
                "section_path": "幻灯片 19 / 幻灯片 19 / body-2",
                "plain_text": "（OCR、语音识别等）",
                "markdown_text": "（OCR、语音识别等）",
            },
            {
                "id": "slide19-model",
                "document_id": "doc-new",
                "version_id": "ver-new",
                "file_name": "【公司介绍】轩辕网络公司介绍202606.pptx",
                "source_type": "bootstrap",
                "trust_level": "internal",
                "page_or_slide": "slide-19",
                "section_path": "幻灯片 19 / 幻灯片 19 / body-1",
                "plain_text": "通用大模型",
                "markdown_text": "通用大模型",
            },
        ]
    )

    reranker = RecordingReranker()
    service = RetrievalService(
        FakeRepository(rows),
        FakeEmbeddingService(),
        reranker,
        FakeVectorStore(),
        candidates=32,
        default_top_k=6,
        retrieval_mode="hybrid",
    )

    result = service.retrieve("基础模型页除了通用大模型，还列了哪些感知或解析能力？", top_k=6)

    assert "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据" in reranker.seen_texts
    top3_texts = [hit.plain_text for hit in result.hits[:3]]
    assert "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据" in top3_texts
    assert "（OCR、语音识别等）" in top3_texts


def test_grounded_stays_false_for_partial_business_architecture_summary_hits():
    hits = [
        build_hit("业务架构：有产懂教，双轮驱动", section_path="slide-11 / body"),
        build_hit("轩辕产教融合建设及运营解决方案", section_path="slide-11 / body"),
        build_hit("双轮驱动，构建核心优势", section_path="slide-15 / body"),
    ]
    hits[0].rerank_score = 7.0
    hits[1].rerank_score = 1.6
    hits[2].rerank_score = -1.3

    assert not RetrievalService._grounded(
        "如果只根据 PPT 内容概括，轩辕网络的业务架构主线是什么？",
        hits,
        ["双轮驱动", "解决方案"],
    )


def test_grounded_accepts_dense_focus_term_enumeration_hits():
    hits = [
        build_hit(
            "AI人才培养与素养提升\n\n产教融合人才培养整体解决方案（1+1+N)\n\n"
            "人才培养服务\n\n师资培养服务\n\n教学资源开发服务\n\n科学研究服务",
            section_path="slide-20 / body",
        ),
        build_hit("1+1+N：（一个模型，双层含义）", section_path="slide-69 / body"),
    ]
    hits[0].rerank_score = 8.7
    hits[1].rerank_score = 3.6

    assert RetrievalService._grounded(
        "产教融合人才培养整体解决方案“1+1+N”包含哪些服务模块？",
        hits,
        ["功能模块", "实训模块", "师资培养服务", "教学资源开发服务", "科学研究服务", "人才培养服务"],
    )


def test_grounded_accepts_enumeration_hits_with_multiple_environment_terms():
    hits = [
        build_hit(
            "AI科学计算平台\n\n智能体应用开发平台\n\n高性能存储资源\n\n高速网络\n\n"
            "AI实训室\n\n产业技术及应用展厅\n\nAIGC赋能中心",
            section_path="slide-21 / body",
        ),
        build_hit("“AI+教育”基础环境\n\n轩辕汇智AI中台", section_path="slide-19 / body"),
    ]
    hits[0].rerank_score = 4.7
    hits[1].rerank_score = 0.8

    assert RetrievalService._grounded(
        "基础环境页除了算力资源，还提到了哪些场地或平台？",
        hits,
        ["基础设施", "平台能力", "存储资源", "高速网络"],
    )


def test_grounded_accepts_four_item_service_enumeration_wording():
    hits = [
        build_hit(
            "产教融合人才培养整体解决方案（1+1+N)\n\n"
            "人才培养服务\n\n师资培养服务\n\n教学资源开发服务\n\n科学研究服务",
            section_path="slide-20 / body",
        ),
        build_hit("人才培养服务", section_path="slide-35 / body"),
    ]
    hits[0].rerank_score = 4.1
    hits[1].rerank_score = 0.8

    assert RetrievalService._grounded(
        "“1+1+N”整体解决方案里，PPT 至少列出的四项服务是什么？",
        hits,
        ["师资培养服务", "教学资源开发服务", "科学研究服务", "人才培养服务"],
    )


def test_ragflow_retrieval_service_maps_chunks():
    def handler(request: httpx.Request) -> httpx.Response:
        assert request.url.path == "/api/v1/retrieval"
        payload = json.loads(request.content.decode("utf-8"))
        assert payload["dataset_ids"] == ["dataset-1"]
        return httpx.Response(
            200,
            json={
                "code": 0,
                "data": {
                    "chunks": [
                        {
                            "id": "chunk-r1",
                            "content": "产业学院采用理事会领导下的院长负责制，理事会是最高决策机构。",
                            "document_id": "doc-r1",
                            "document_name": "产业学院资料.pptx",
                            "dataset_id": "dataset-1",
                            "positions": ["slide 12"],
                            "similarity": 0.91,
                            "vector_similarity": 0.82,
                            "term_similarity": 0.95,
                        }
                    ],
                    "total": 1,
                },
            },
        )

    transport = httpx.MockTransport(handler)
    http_client = httpx.Client(base_url="http://ragflow.local", transport=transport)
    client = RagflowClient(base_url="http://ragflow.local", api_key="token", client=http_client)
    service = RagflowRetrievalService(
        client=client,
        dataset_ids=["dataset-1"],
        document_ids=[],
        default_top_k=3,
    )

    result = service.retrieve("产业学院采用什么治理模式？", top_k=2, focus_terms=["治理模式"])

    assert result.hits
    assert result.hits[0].file_name == "产业学院资料.pptx"
    assert result.hits[0].page_or_slide == "slide 12"
    assert result.hits[0].source_type == "ragflow"
    assert result.hits[0].raw_scores["dataset_id"] == "dataset-1"
    assert result.grounded is True
    http_client.close()


def test_ragflow_retrieval_service_normalizes_array_like_positions():
    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(
            200,
            json={
                "code": 0,
                "data": {
                    "chunks": [
                        {
                            "id": "chunk-r1",
                            "content": "核心定位：紧扣根技术筑基、产教融育人、师范践初心主线。",
                            "document_id": "doc-r1",
                            "document_name": "展厅资料.pptx",
                            "dataset_id": "dataset-1",
                            "positions": ["[1, 0, 0, 0, 0]"],
                            "similarity": 0.91,
                            "vector_similarity": 0.82,
                            "term_similarity": 0.95,
                        }
                    ],
                    "total": 1,
                },
            },
        )

    transport = httpx.MockTransport(handler)
    http_client = httpx.Client(base_url="http://ragflow.local", transport=transport)
    client = RagflowClient(base_url="http://ragflow.local", api_key="token", client=http_client)
    service = RagflowRetrievalService(
        client=client,
        dataset_ids=["dataset-1"],
        document_ids=[],
        default_top_k=3,
    )

    result = service.retrieve("核心定位主线是什么？", top_k=1)

    assert result.hits[0].page_or_slide == "1 > 0 > 0 > 0 > 0"
    assert result.hits[0].section_path == "1 > 0 > 0 > 0 > 0"
    http_client.close()


def test_adaptive_retrieval_prefers_grounded_local_result():
    class FakeService:
        def __init__(self, result: RetrievalResult) -> None:
            self.result = result
            self.calls = 0

        def retrieve(self, question: str, top_k=None, focus_terms=None, expansion_terms=None) -> RetrievalResult:
            self.calls += 1
            return self.result

    local_result = RetrievalResult(
        hits=[build_hit("本地细粒度命中，包含理论重构、架构重构、软件重构", section_path="slide-5 / 文本 1")],
        grounded=True,
        focus_terms=["理论重构"],
        expanded_query="q",
        expansion_terms=[],
    )
    local_result.hits[0].rerank_score = 0.9
    local = FakeService(local_result)

    remote_result = RetrievalResult(
        hits=[build_hit("RAGFlow 粗粒度命中", section_path="5 > 0 > 0 > 0 > 0")],
        grounded=True,
        focus_terms=["理论重构"],
        expanded_query="q",
        expansion_terms=[],
    )
    remote_result.hits[0].source_type = "ragflow"
    remote_result.hits[0].rerank_score = 0.92
    remote = FakeService(remote_result)

    service = AdaptiveRetrievalService(local, remote, local_grounded_score_threshold=0.55)
    result = service.retrieve("华为通过哪三个重构，哪五大方向突围")

    assert result.hits[0].section_path == "slide-5 / 文本 1"
    assert result.backend_path == "local"
    assert result.route_reason == "local_grounded_above_threshold"
    assert result.remote_attempted is False
    assert result.local_top_score == 0.9
    assert result.local_grounded_score_threshold == 0.55
    assert local.calls == 1
    assert remote.calls == 0


def test_adaptive_retrieval_uses_remote_when_local_not_grounded():
    class FakeService:
        def __init__(self, result: RetrievalResult) -> None:
            self.result = result

        def retrieve(self, question: str, top_k=None, focus_terms=None, expansion_terms=None) -> RetrievalResult:
            return self.result

    local_result = RetrievalResult(
        hits=[build_hit("本地命中较弱", section_path="docx / A")],
        grounded=False,
        focus_terms=[],
        expanded_query="q",
        expansion_terms=[],
    )
    local_result.hits[0].rerank_score = 0.2
    remote_result = RetrievalResult(
        hits=[build_hit("远端命中更完整，包含功能全面、性能优异、全球共享", section_path="5 > 643 > 957")],
        grounded=True,
        focus_terms=[],
        expanded_query="q",
        expansion_terms=[],
    )
    remote_result.hits[0].source_type = "ragflow"
    remote_result.hits[0].rerank_score = 0.82

    service = AdaptiveRetrievalService(FakeService(local_result), FakeService(remote_result))
    result = service.retrieve("华为人才的优势是什么")

    assert result.hits[0].source_type == "ragflow"
    assert result.backend_path == "ragflow"
    assert result.used_fallback is True
    assert result.fallback_reason == "remote_selected_after_local_insufficient"
    assert result.route_reason == "remote_grounded_local_not_grounded"
    assert result.remote_attempted is True
    assert result.local_top_score == 0.2
    assert result.local_quality_score is not None
    assert result.remote_quality_score is not None


def test_ragflow_retrieval_service_retries_without_keyword_when_chat_model_missing():
    calls: list[dict] = []

    def handler(request: httpx.Request) -> httpx.Response:
        payload = json.loads(request.content.decode("utf-8"))
        calls.append(payload)
        if len(calls) == 1:
            return httpx.Response(200, json={"code": 100, "message": "Exception('No default chat model is set.')"})
        return httpx.Response(
            200,
            json={
                "code": 0,
                "data": {
                    "chunks": [
                        {
                            "id": "chunk-r2",
                            "content": "核心定位：紧扣根技术筑基、产教融育人、师范践初心主线。",
                            "document_id": "doc-r2",
                            "document_name": "展厅资料.pptx",
                            "dataset_id": "dataset-1",
                            "positions": ["slide 1"],
                            "similarity": 0.93,
                            "vector_similarity": 0.85,
                            "term_similarity": 0.0,
                        }
                    ],
                    "total": 1,
                },
            },
        )

    transport = httpx.MockTransport(handler)
    http_client = httpx.Client(base_url="http://ragflow.local", transport=transport)
    client = RagflowClient(base_url="http://ragflow.local", api_key="token", client=http_client)
    service = RagflowRetrievalService(
        client=client,
        dataset_ids=["dataset-1"],
        document_ids=[],
        default_top_k=3,
        keyword=True,
    )

    result = service.retrieve("核心定位主线是什么？", top_k=1)

    assert result.hits[0].file_name == "展厅资料.pptx"
    assert calls[0]["keyword"] is True
    assert calls[1]["keyword"] is False
    assert calls[1]["use_kg"] is False
    assert calls[1]["toc_enhance"] is False
    http_client.close()


def test_fallback_retrieval_service_uses_local_when_primary_fails():
    class BrokenPrimary:
        def retrieve(self, question, top_k=None, focus_terms=None, expansion_terms=None):
            raise RuntimeError("ragflow unavailable")

    class WorkingFallback:
        def retrieve(self, question, top_k=None, focus_terms=None, expansion_terms=None):
            return "fallback-result"

    service = FallbackRetrievalService(BrokenPrimary(), WorkingFallback())
    assert service.retrieve("测试问题") == "fallback-result"


def test_fallback_retrieval_service_uses_local_when_primary_times_out():
    import time

    class SlowPrimary:
        def retrieve(self, question, top_k=None, focus_terms=None, expansion_terms=None):
            time.sleep(0.05)
            return "primary-result"

    class WorkingFallback:
        def retrieve(self, question, top_k=None, focus_terms=None, expansion_terms=None):
            return RetrievalResult(
                hits=[],
                grounded=False,
                focus_terms=[],
                expanded_query=question,
                expansion_terms=[],
            )

    service = FallbackRetrievalService(SlowPrimary(), WorkingFallback(), primary_timeout_ms=5)
    result = service.retrieve("测试问题")
    assert isinstance(result, RetrievalResult)
    assert result.used_fallback is True
    assert result.fallback_reason == "primary_timeout"


def test_llm_context_block_uses_compact_evidence_only():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit(
            "产业学院治理结构方面采用理事会领导下的院长负责制。理事会作为最高决策机构。决策会议每季度召开1次。",
            section_path="治理结构",
        )
    ]
    context = service._context_block(
        "产业学院采用什么治理模式？",
        "factoid",
        ["治理模式", "产业学院"],
        citations,
    )
    assert "理事会领导下的院长负责制" in context
    assert "kb.docx" not in context
    assert "文件:" not in context
    assert "位置:" not in context
    assert "信任级别:" not in context


def test_teaching_materials_answer_uses_material_list():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit("教学资料包括教学大纲、MOOC、授课PPT、电子教材、实验手册、实验室搭建指南。"),
    ]
    answer, grounded = service._compose_extract_answer(
        "华为ICT学院的教学资料包括哪些内容？",
        "enumeration",
        "教学资料",
        ["教学资料", "华为ICT学院"],
        citations,
        True,
    )
    assert "教学大纲" in answer
    assert "MOOC" in answer
    assert "授课PPT" in answer
    assert "通识课" not in answer
    assert "教学大纲" in grounded


def test_source_leak_detector_does_not_flag_page_words():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    assert not service._contains_source_leak("进入首页左下角即可查看设备IP地址。", [])
    assert not service._contains_source_leak("在页面左下角可以看到设备编号。", [])
    assert service._contains_source_leak("详见第3页。", [])


def test_question_type_and_followup_focus_rules():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    assert service._infer_question_type("在哪里查看边缘网关的IP地址？") == "procedure"
    assert service._infer_question_type("Edge智控APP绑定失败时应该先检查什么？") == "procedure"
    assert service._answer_matches_focus(
        "理事会每季度召开1次决策会议。",
        "产业学院决策会议的召开频率",
        ["决策会议", "召开频率"],
        "那决策会议多久召开一次？",
    )
    assert service._answer_matches_focus(
        "实训套件采用端、边、云、应用四层架构设计。",
        "华为AR502H系列工业级边缘计算网关采用什么技术架构？",
        ["技术架构", "AR502H"],
        "它采用什么技术架构？",
    )
    assert service._answer_matches_focus(
        "最高决策机构是产业学院理事会。",
        "产业学院的最高决策机构名称",
        ["最高决策机构", "理事会"],
        "它的最高决策机构是什么？",
    )
    assert service._answer_matches_focus(
        "支持本地部署DeepSeek、Qwen等开源大模型。",
        "协作式机械臂产品支持本地部署的大模型",
        ["本地部署", "大模型", "DeepSeek", "Qwen"],
        "它支持本地部署哪些大模型？",
    )
    assert service._answer_matches_focus(
        "核心标语是“根生万物·智育未来”。",
        "根技术体验中心的核心标语",
        ["根技术体验中心的核心标语"],
        "根技术体验中心的核心标语是什么？",
    )
    assert service._answer_matches_focus(
        "三位一体方向是根技术、人工智能、职教母机。",
        "产业学院强调的三位一体方向",
        ["产业学院强调的三位一体方向"],
        "产业学院强调的三位一体方向是什么？",
    )
    assert service._answer_matches_focus(
        "产品采用两台协作机器人和两套视觉系统。",
        "协作式机械臂产品几台协作机器人和几套视觉系统",
        ["协作式机械臂产品", "几台协作机器人和几套视觉系统"],
        "协作式机械臂产品采用几台协作机器人和几套视觉系统？",
    )


def test_enumeration_and_insufficient_guards():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    assert service._enumeration_answer(
        [
            "沟通机制：建立“月例会、季汇报、年总结”制度，同步进展、协调问题。",
            "决策机制：理事会每季度召开1次决策会议。",
        ]
    ) == "包括月例会、季汇报、年总结"
    assert service._signals_insufficient_text("当前知识库没有直接证据。现有资料未提及学费信息。")


def test_special_case_answer_and_review_issue_pruning():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit(
            "1. 决策机制：理事会每季度召开 1 次决策会议；"
            "2. 执行机制：院长牵头；"
            "3. 沟通机制：建立“月例会、季汇报、年总结”制度，同步进展、协调问题。",
            section_path="机制保障 / 协同机制",
        )
    ]
    answer, grounded = service._compose_extract_answer(
        "产业学院有哪些沟通制度？",
        "enumeration",
        "产业学院、沟通制度",
        ["产业学院", "沟通制度", "月例会", "季汇报", "年总结"],
        citations,
        True,
    )
    assert answer == "包括月例会、季汇报、年总结。"
    assert "月例会" in grounded

    env_answer, env_grounded = service._compose_extract_answer(
        "协作式机械臂产品的开放性实验环境是什么？",
        "factoid",
        "协作式机械臂产品的开放性实验环境",
        ["开放性实验环境", "Jupyter Notebook"],
        [
            build_hit(
                "实验代码在Jupyter Notebook环境中编写，具有如下功能："
                "教师与学生可以在浏览器上直接进行交互式编程实验。",
                section_path="三、功能与应用场景 / 1.开放性实验环境",
            )
        ],
        True,
    )
    assert env_answer == "开放性实验环境主要基于Jupyter Notebook环境。"
    assert "Jupyter Notebook" in env_grounded

    model_answer, model_grounded = service._compose_extract_answer(
        "它支持本地部署哪些大模型？",
        "followup",
        "协作式机械臂产品支持本地部署的大模型",
        ["本地部署", "大模型", "DeepSeek", "Qwen"],
        [
            build_hit(
                "除以上功能外，还在运算单元中完成了DeepSeek、Qwen等开源大模型的本地化部署，"
                "支持大模型技术开发、大模型+视觉、大模型+语音、大模型+机器人等多个行业场景的应用实践。",
                section_path="产品能力 / 大模型部署",
            )
        ],
        True,
    )
    assert model_answer == "支持本地部署DeepSeek、Qwen等开源大模型。"
    assert "DeepSeek" in model_grounded and "Qwen" in model_grounded

    analysis = QueryAnalysis(
        rewritten_query="协作式机械臂产品支持本地部署的大模型有哪些？",
        question_type="followup",
        answer_focus="协作式机械臂产品支持本地部署的大模型",
        focus_terms=["本地部署", "大模型", "DeepSeek", "Qwen"],
    )
    draft = DraftAnswer(
        answer=model_answer,
        grounded_answer=model_grounded,
        inference_note="",
        question_type="followup",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    pruned = service._prune_review_issues(
        ["off_target", "followup_error"],
        model_answer,
        [],
        draft,
        "它支持本地部署哪些大模型？",
        analysis,
    )
    assert pruned == []


def test_compose_extract_answer_handles_procedure_short_answers():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)

    bind_answer, bind_grounded = service._compose_extract_answer(
        "Edge智控APP绑定失败时应该先检查什么？",
        "procedure",
        "绑定失败检查项",
        ["绑定失败", "互联网", "网络认证"],
        [build_hit("如果绑定失败，请先检查实训室网络是否能正常连接互联网；如果访问互联网需要认证，请联系学校网络管理员。")],
        True,
    )
    assert "互联网" in bind_answer
    assert "网络管理员" in bind_answer
    assert "互联网" in bind_grounded

    ip_answer, ip_grounded = service._compose_extract_answer(
        "在哪里查看边缘网关的IP地址？",
        "procedure",
        "查看IP地址位置",
        ["IP地址", "首页左下角", "Edge智控"],
        [build_hit("打开Edge智控APP进入首页后，可在左下角查看设备IP地址。")],
        True,
    )
    assert "首页左下角" in ip_answer
    assert "左下角" in ip_grounded

    focus_answer, focus_grounded = service._compose_extract_answer(
        "摄像头出现虚焦时该怎么处理？",
        "procedure",
        "摄像头虚焦处理",
        ["虚焦", "旋转镜头", "固定螺丝"],
        [build_hit("松开固定螺丝后左右旋转镜头调焦，调到合适位置后再拧紧。")],
        True,
    )
    assert "松开固定螺丝" in focus_answer
    assert "旋转镜头" in focus_answer
    assert "拧紧" in focus_answer
    assert "拧紧" in focus_grounded

    password_answer, password_grounded = service._compose_extract_answer(
        "如果忘记设备登录密码，应该怎么处理？",
        "procedure",
        "忘记设备登录密码处理方式",
        ["登录密码", "产品手册", "账号密码"],
        [build_hit("如果忘记设备登录密码，可查询对应产品手册，里面提供了所需的账号和密码信息。")],
        True,
    )
    assert "产品手册" in password_answer
    assert "账号和密码" in password_grounded
    assert "账号密码" in password_answer

    scale_answer, scale_grounded = service._compose_extract_answer(
        "电子秤称重不准确时应该怎么处理？",
        "procedure",
        "电子秤称重不准确处理方式",
        ["电子秤", "称重", "清零"],
        [build_hit("电子秤开机状态下，电子秤圆盘上不放置任何物件，按电子秤屏幕下方的清零按钮，可对电子秤执行清零操作后，再进行称重。")],
        True,
    )
    assert "清零按钮" in scale_answer
    assert "重新称重" in scale_answer
    assert "不放置任何物件" in scale_grounded


def test_compose_extract_answer_handles_compact_factoid_and_enumeration_answers():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)

    slogan_answer, slogan_grounded = service._compose_extract_answer(
        "根技术体验中心的核心标语是什么？",
        "factoid",
        "根技术体验中心核心标语",
        ["核心标语", "根生万物", "智育未来"],
        [build_hit("体验中心以“根生万物·智育未来”为核心标语。")],
        True,
    )
    assert slogan_answer == "核心标语是“根生万物·智育未来”。"
    assert "根生万物" in slogan_grounded

    pillars_answer, pillars_grounded = service._compose_extract_answer(
        "根技术体验中心聚焦的四个支柱方向有哪些？",
        "enumeration",
        "四个支柱方向",
        ["四个支柱", "智慧农业", "智能制造", "健康卫生", "智能教育"],
        [build_hit("呼应四个支柱：智慧农业、智能制造、生命健康、智能教育。")],
        True,
    )
    assert pillars_answer == "包括智慧农业、智能制造、健康卫生、智能教育。"
    assert "智慧农业" in pillars_grounded

    mainline_answer, mainline_grounded = service._compose_extract_answer(
        "华为根技术体验中心的核心定位主线是什么？",
        "factoid",
        "核心定位主线",
        ["核心定位主线", "根技术筑基", "产教融育人", "师范践初心"],
        [build_hit("核心定位主线：根技术筑基、产教融育人、师范践初心。", section_path="展厅定位")],
        True,
    )
    assert mainline_answer == "核心定位主线是“根技术筑基、产教融育人、师范践初心”。"
    assert "根技术筑基" in mainline_grounded

    vision_answer, vision_grounded = service._compose_extract_answer(
        "协作式机械臂的视觉应用包含哪些方向？",
        "enumeration",
        "视觉应用方向",
        ["视觉应用", "定位", "检测", "识别"],
        [build_hit("视觉应用包括定位、检测、识别三个方向。", section_path="视觉应用")],
        True,
    )
    assert vision_answer == "包括定位、检测、识别。"
    assert "定位" in vision_grounded and "识别" in vision_grounded

    modules_answer, modules_grounded = service._compose_extract_answer(
        "协作式机械臂产品包含哪些模块？",
        "enumeration",
        "协作式机械臂产品模块",
        ["模块", "仓储模块", "视觉识别与分拣模块", "语音交互模块"],
        [build_hit("功能模块包括仓储模块、视觉识别与分拣模块、语音交互模块。")],
        True,
    )
    assert modules_answer == "包括仓储模块、视觉识别与分拣模块、语音交互模块。"
    assert "语音交互模块" in modules_grounded

    certifications_answer, certifications_grounded = service._compose_extract_answer(
        "产业学院的认证覆盖哪些级别？",
        "enumeration",
        "产业学院认证级别",
        ["认证", "HCIA", "HCIP", "HCIE"],
        [build_hit("认证培训课程覆盖HCIA、HCIP、HCIE等证书课程。")],
        True,
    )
    assert certifications_answer == "包括HCIA、HCIP、HCIE。"
    assert "HCIP" in certifications_grounded


def test_compose_extract_answer_supports_company_ppt_factoid_and_enumeration_boundaries():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)

    company_answer, company_grounded = service._compose_extract_answer(
        "公司概况页提到轩辕网络深耕教育多少年、专注什么方向？",
        "factoid",
        "公司概况、轩辕网络",
        ["公司概况", "轩辕网络", "深耕教育", "专注", "产教融合"],
        [build_hit("01 公司概况 28 年教育深耕者 专注产教融合 · 成就职业教育未来", section_path="幻灯片 3 / body")],
        True,
    )
    assert company_answer == "轩辕网络深耕教育28年，专注产教融合方向。"
    assert "教育深耕者" in company_grounded

    services_answer, services_grounded = service._compose_extract_answer(
        "产教融合人才培养整体解决方案“1+1+N”包含哪些服务模块？",
        "enumeration",
        "产教融合、人才培养",
        ["1+1+N", "服务模块", "人才培养服务", "师资培养服务", "教学资源开发服务", "科学研究服务"],
        [
            build_hit("AI人才培养与素养提升 产教融合人才培养整体解决方案（1+1+N) 构建“政-行-企-校” 多元协同育人新生态", section_path="幻灯片 20 / body"),
            build_hit("高等院校 人才培养服务 师资培养服务 教学资源开发服务 科学研究服务 计算机类专业 产教融合运营平台", section_path="幻灯片 20 / body"),
        ],
        True,
    )
    assert services_answer == "包括人才培养服务、师资培养服务、教学资源开发服务、科学研究服务。"
    assert "科学研究服务" in services_grounded

    env_answer, env_grounded = service._compose_extract_answer(
        "基础环境页除了算力资源，还提到了哪些场地或平台？",
        "enumeration",
        "基础环境、场地平台",
        ["基础环境", "场地", "平台", "数智技术实践中心", "产业技术及应用展厅", "AIGC实战平台"],
        [
            build_hit("基础环境 数智技术实践中心 通用算力资源 智能算力资源 高性能存储资源 高速网络", section_path="幻灯片 18 / body"),
            build_hit("产业技术及应用展厅 AIGC实战平台 AIGC赋能中心", section_path="幻灯片 21 / body"),
        ],
        True,
    )
    assert env_answer == "包括数智技术实践中心、产业技术及应用展厅、AIGC实战平台、AIGC赋能中心。"
    assert "产业技术及应用展厅" in env_grounded

    positioning_answer, positioning_grounded = service._compose_extract_answer(
        "战略定位页有没有把轩辕网络定义成“AI+产教融合服务商”？",
        "factoid",
        "战略定位、轩辕网络",
        ["战略定位", "轩辕网络", "AI+产教融合服务商"],
        [
            build_hit("战略定位：AI+产教融合服务商", section_path="幻灯片 16 / body"),
            build_hit("广东轩辕网络科技股份有限公司是业内领先的AI+产教融合服务商。", section_path="幻灯片 4 / body-1"),
        ],
        True,
    )
    assert positioning_answer == "是，战略定位页把轩辕网络定义为AI+产教融合服务商。"
    assert "AI+产教融合服务商" in positioning_grounded


def test_compose_extract_answer_handles_definition_goal_and_curriculum_factoids():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)

    definition_answer, definition_grounded = service._compose_extract_answer(
        "华为根技术是什么？",
        "factoid",
        "华为根技术定义",
        ["华为根技术", "定义"],
        [build_hit("总结 华为根技术指的是华为在信息通信技术（ICT）领域自主研发、长期积累、持续投入的核心底层技术。")],
        True,
    )
    assert definition_answer == "华为根技术指的是华为在信息通信技术（ICT）领域自主研发、长期积累、持续投入的核心底层技术。"
    assert "核心底层技术" in definition_grounded

    goal_answer, goal_grounded = service._compose_extract_answer(
        "广师大根技术体验中心的建设目标是什么？",
        "factoid",
        "建设目标",
        ["建设目标", "根技术筑基", "师范践初心"],
        [build_hit("建设目标：根技术筑基、产教融育人、师范践初心。")],
        True,
    )
    assert goal_answer == "建设目标是根技术筑基、产教融育人、师范践初心。"
    assert "根技术筑基" in goal_grounded

    curriculum_answer, curriculum_grounded = service._compose_extract_answer(
        "根技术课程体系",
        "factoid",
        "根技术课程体系",
        ["课程体系", "17个学院70个专业", "新师范"],
        [
            build_hit(
                "介绍广师大根技术通识教育课程体系（17个学院70个专业，针对不同学科门类特点，打造新师范、新工科、新文科）。"
            ),
            build_hit("在线学习模块：可查看通识课课件、AIGC实战平台，适配学生自主学习、访客快速了解根技术的需求。"),
        ],
        True,
    )
    assert "17个学院70个专业" in curriculum_answer
    assert "新师范、新工科、新文科" in curriculum_answer
    assert "AIGC实战平台" in curriculum_answer
    assert "课程体系" in curriculum_grounded

    g1_answer, g1_grounded = service._compose_extract_answer(
        "宇树G1机器人的关节数量是多少？",
        "factoid",
        "宇树G1关节数量",
        ["宇树G1", "关节数量", "总自由度"],
        [
            build_hit(
                "具身智能机器人：多场景自主作业 身型数值 总自由度 s43个 体重约35kg Unitree G1 身高约130cm。",
                section_path="具身智能实训区",
            )
        ],
        True,
    )
    assert g1_answer == "Unitree G1的总自由度约为43个。"
    assert "43个" in g1_grounded


def test_pattern_based_factoid_prefers_mainline_over_generic_definition():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)

    answer = service._pattern_based_factoid_answer(
        "华为根技术体验中心的核心定位主线是什么？",
        [
            build_hit("核心定位：紧扣“根技术筑基、产教融育人、师范践初心”主线。", section_path="展厅定位"),
            build_hit("华为根技术指的是华为在信息通信技术领域自主研发、长期积累、持续投入的核心底层技术。", section_path="定义"),
        ],
    )

    assert answer == "核心定位主线是“根技术筑基、产教融育人、师范践初心”。"


def test_pattern_based_factoid_supports_hall_paraphrases():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)

    mainline = service._pattern_based_factoid_answer(
        "展厅文化建设紧扣的是哪三句话？",
        [build_hit("文化建设核心：紧扣“根技术筑基、产教融育人、师范践初心”主线。", section_path="展厅定位")],
    )
    slogan = service._pattern_based_factoid_answer(
        "这个展厅的口号是什么？",
        [build_hit("建设思路：以“根生万物·智育未来”为核心标语，凸显校园特色与展厅定位。", section_path="LED落地造型墙")],
    )

    assert mainline == "核心定位主线是“根技术筑基、产教融育人、师范践初心”。"
    assert slogan == "核心标语是“根生万物·智育未来”。"


def test_pattern_based_factoid_supports_architecture_course_meeting_and_load():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)

    architecture = service._pattern_based_factoid_answer(
        "产业学院采用什么技术应用架构？",
        [build_hit("采用“底座 + 支柱”技术应用架构，全面保障产业学院高质量发展。")],
    )
    assert architecture == "技术应用架构采用“底座 + 支柱”架构。"

    course = service._pattern_based_factoid_answer(
        "产业学院提到的AI核心课程是什么？",
        [build_hit("深化协同联动，共建《现代教育技术与智慧教学》AI 赋能核心课程。")],
    )
    assert "现代教育技术与智慧教学" in course

    meeting = service._pattern_based_factoid_answer(
        "产业学院的决策会议多久召开一次？",
        [build_hit("决策机制：理事会每季度召开 1 次决策会议，审定发展战略。")],
    )
    assert meeting == "决策会议每季度召开1次。"

    load = service._pattern_based_factoid_answer(
        "协作机器人的额定负载是多少？",
        [build_hit("<table><tr><td>额定负载</td><td>3kg</td></tr></table>")],
    )
    assert load == "协作机器人的额定负载是3kg。"

    layout = service._pattern_based_factoid_answer(
        "华为在根技术研发布局是什么？",
        [build_hit("华为在根技术研发布局：强力投入研究与开发，以创新驱动未来发展。")],
    )
    assert layout == "华为在根技术研发布局上强调强力投入研究与开发，以创新驱动未来发展。"

    breakthroughs = service._pattern_based_factoid_answer(
        "华为通过哪三个重构，哪五大方向突围？",
        [build_hit("华为通过3个重构，5大方向突围。理论重构、架构重构、软件重构。基础理论、基础硬件、基础软件、开发工具、运营系统。")],
    )
    assert "理论重构、架构重构、软件重构" in breakthroughs

    intro = service._pattern_based_factoid_answer(
        "给我介绍一下华为ICT学院",
        [build_hit("华为ICT学院是华为主导的、面向全球的校企合作项目。")],
    )
    assert "校企合作项目" in intro

    advantage = service._pattern_based_factoid_answer(
        "华为人才的优势是什么？",
        [build_hit("华为人才在线官网优势：功能全面、性能优异、全球共享、操作灵活、效果评价。")],
    )
    assert "功能全面" in advantage


def test_procedure_short_answer_supports_huawei_ict_application():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    answer, grounded = service._procedure_short_answer("怎么样提交申请？")
    assert "了解华为ICT学院项目内容及要求" in answer
    assert "华为审核" in grounded


def test_grounded_supports_huawei_root_layout_and_application_questions():
    layout_hits = [build_hit("华为在根技术研发布局：强力投入研究与开发，以创新驱动未来发展。", section_path="幻灯片 4")]
    breakthrough_hits = [build_hit("华为通过3个重构，5大方向突围。理论重构、架构重构、软件重构。基础理论、基础硬件、基础软件、开发工具、运营系统。", section_path="幻灯片 5")]
    apply_hits = [build_hit("华为ICT学院申请步骤：在充分了解项目内容及要求后，填写注册信息、提交相关申请，由华为审核并通知结果。", section_path="申请指南")]

    assert RetrievalService._grounded("华为在根技术研发布局是什么？", layout_hits, ["强力投入研究与开发", "创新驱动未来发展"])
    assert RetrievalService._grounded("华为通过哪三个重构，哪五大方向突围？", breakthrough_hits, ["理论重构", "运营系统"])
    assert RetrievalService._grounded("怎么样提交申请？", apply_hits, ["申请步骤", "提交相关申请", "华为审核"])


def test_grounded_supports_hall_and_exhibit_paraphrases():
    mainline_hits = [
        build_hit("文化建设核心：紧扣“根技术筑基、产教融育人、师范践初心”主线。", section_path="展厅定位")
    ]
    slogan_hits = [
        build_hit("建设思路：以“根生万物·智育未来”为核心标语，凸显校园特色与展厅定位。", section_path="LED落地造型墙")
    ]
    exhibit_hits = [
        build_hit("展品：鸿蒙智联场景应用实训箱，Atlas智能小车。", section_path="鸿蒙万物互联体验区")
    ]

    assert RetrievalService._grounded("展厅文化建设紧扣的是哪三句话？", mainline_hits, ["核心定位主线", "根技术筑基"])
    assert RetrievalService._grounded("这个展厅的口号是什么？", slogan_hits, ["核心标语", "根生万物"])
    assert RetrievalService._grounded(
        "鸿蒙智能装备体验区展示了哪些设备？",
        exhibit_hits,
        ["鸿蒙智联场景应用实训箱", "Atlas智能小车"],
    )


def test_compose_extract_answer_prefers_arm_courses_over_majors():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    answer, grounded = service._compose_extract_answer(
        "协作式机械臂产品适用哪些课程？",
        "enumeration",
        "协作式机械臂适用课程",
        ["适用课程", "Python程序设计", "深度学习"],
        [
            build_hit(
                "面向专业：人工智能、机器人工程、智能制造、自动化、电子、信息科学、机电等专业。"
                "适用课程：Python程序设计、深度学习、数字图像处理、机器视觉、基于视觉的机器人应用、大模型技术应用。",
                section_path="面向专业和课程",
            ),
            build_hit("教学科研实验：机器视觉识别、机器人控制、AI算法优化验证。", section_path="机器视觉实训区"),
        ],
        True,
    )
    assert answer == "包括Python程序设计、深度学习、数字图像处理、机器视觉、基于视觉的机器人应用等。"
    assert "大模型技术应用" in grounded


def test_deterministic_review_flags_overlong_factoid_draft():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [build_hit("产业学院采用理事会领导下的院长负责制。", section_path="治理结构")]
    analysis = QueryAnalysis(
        rewritten_query="产业学院采用什么治理模式？",
        question_type="factoid",
        answer_focus="产业学院治理模式",
        focus_terms=["治理模式", "理事会领导下的院长负责制"],
    )
    draft = DraftAnswer(
        answer="产业学院采用基于理事会的治理模式，具体实行理事会领导下的院长负责制。该模式强调科学治理、协同推进和高效运作。",
        grounded_answer="产业学院采用理事会领导下的院长负责制。",
        inference_note="",
        question_type="factoid",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    review = service._deterministic_review("产业学院采用什么治理模式？", analysis, citations, draft)
    assert "verbose" in review.issues
    assert "direct" in review.issues
    assert review.revised_answer == "产业学院采用理事会领导下的院长负责制。"


def test_deterministic_review_blocks_partial_business_architecture_summary_release():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit(
            "业务架构：有产懂教，双轮驱动。轩辕产教融合建设及运营解决方案。",
            section_path="幻灯片 11 / body",
        ),
        build_hit(
            "广东轩辕网络科技股份有限公司是业内领先的AI+产教融合服务商。",
            section_path="幻灯片 4 / body-1",
        ),
    ]
    analysis = QueryAnalysis(
        rewritten_query="轩辕网络的业务架构如何概括？",
        question_type="factoid",
        answer_focus="业务架构概括",
        focus_terms=["业务架构", "双轮驱动", "产教融合建设及运营解决方案"],
    )
    draft = DraftAnswer(
        answer="轩辕网络的业务架构可以概括为有产懂教、双轮驱动，是业内领先的AI+产教融合服务商。",
        grounded_answer="业务架构：有产懂教，双轮驱动。轩辕产教融合建设及运营解决方案。",
        inference_note="",
        question_type="factoid",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )

    review = service._deterministic_review("轩辕网络的业务架构如何概括？", analysis, citations, draft)
    payload = service.finalize_answer("轩辕网络的业务架构如何概括？", analysis, citations, draft, review)

    assert "unsupported" in review.issues
    assert payload["grounded"] is False
    assert payload["answer"] == "当前知识库中没有找到相关信息。"


def test_review_answer_keeps_business_architecture_unsupported_issue():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit(
            "业务架构：有产懂教，双轮驱动。轩辕产教融合建设及运营解决方案。",
            section_path="幻灯片 11 / body",
        ),
        build_hit(
            "广东轩辕网络科技股份有限公司是业内领先的AI+产教融合服务商。",
            section_path="幻灯片 4 / body-1",
        ),
    ]
    analysis = QueryAnalysis(
        rewritten_query="轩辕网络的业务架构在 PPT 中是如何概括的？",
        question_type="factoid",
        answer_focus="业务架构",
        focus_terms=["业务架构", "双轮驱动"],
    )
    draft = DraftAnswer(
        answer=(
            "1. 广东轩辕网络科技股份有限公司是业内领先的AI+产教融合服务商。 "
            "2. 2026年6月 | 中国·广州 数智人才共育，教育产业共赢 轩辕网络公司介绍 2026年6月 "
            "3. 由轩辕网络与广铁共建产业学院，轩辕网络负责产业学院一体化运营服务。"
        ),
        grounded_answer="业务架构：有产懂教，双轮驱动。轩辕产教融合建设及运营解决方案。",
        inference_note="",
        question_type="factoid",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )

    review = service.review_answer("轩辕网络的业务架构在 PPT 中是如何概括的？", analysis, citations, draft)
    payload = service.finalize_answer("轩辕网络的业务架构在 PPT 中是如何概括的？", analysis, citations, draft, review)

    assert "unsupported" in review.issues
    assert review.reviewer_intervened is True
    assert payload["grounded"] is False
    assert payload["answer"] == "当前知识库中没有找到相关信息。"


def test_deterministic_review_blocks_foundation_capability_partial_enumeration_release():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit("（OCR、语音识别等）", section_path="幻灯片 19 / body-2"),
        build_hit("训练专属大模型。", section_path="幻灯片 19 / body-1"),
    ]
    analysis = QueryAnalysis(
        rewritten_query="基础模型页除了通用大模型，还列了哪些感知或解析能力？",
        question_type="enumeration",
        answer_focus="基础模型感知或解析能力",
        focus_terms=["基础模型", "OCR", "语音识别", "文档增强解析", "知识元数据"],
    )
    draft = DraftAnswer(
        answer="包括OCR、语音识别等。",
        grounded_answer="（OCR、语音识别等）",
        inference_note="",
        question_type="enumeration",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )

    review = service._deterministic_review(
        "基础模型页除了通用大模型，还列了哪些感知或解析能力？",
        analysis,
        citations,
        draft,
    )
    payload = service.finalize_answer(
        "基础模型页除了通用大模型，还列了哪些感知或解析能力？",
        analysis,
        citations,
        draft,
        review,
    )

    assert "unsupported" in review.issues
    assert payload["grounded"] is False
    assert payload["answer"] == "当前知识库中没有找到相关信息"


def test_finalize_blocks_foundation_capability_partial_enumeration_even_without_review_issue():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit("（OCR、语音识别等）", section_path="幻灯片 19 / body-2"),
        build_hit("训练专属大模型。", section_path="幻灯片 19 / body-1"),
    ]
    analysis = QueryAnalysis(
        rewritten_query="基础模型页除了通用大模型，还列了哪些感知或解析能力？",
        question_type="enumeration",
        answer_focus="基础模型感知或解析能力",
        focus_terms=["基础模型", "OCR", "语音识别", "文档增强解析", "知识元数据"],
    )
    draft = DraftAnswer(
        answer="包括OCR、语音识别等。",
        grounded_answer="（OCR、语音识别等）",
        inference_note="",
        question_type="enumeration",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    review = ReviewResult(
        passed=True,
        issues=[],
        revised_answer=draft.answer,
        revised_grounded_answer=draft.grounded_answer,
        revised_inference_note="",
        risk_level="low",
        reviewer_intervened=False,
    )

    payload = service.finalize_answer(
        "基础模型页除了通用大模型，还列了哪些感知或解析能力？",
        analysis,
        citations,
        draft,
        review,
    )

    assert payload["grounded"] is False
    assert payload["answer"] == "当前知识库中没有找到相关信息"


def test_finalize_blocks_foundation_platform_capability_partial_coverage_release():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit(
            "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
            section_path="幻灯片 19 / body",
        ),
        build_hit("（deepseek、通义千问）", section_path="幻灯片 19 / body-1"),
        build_hit("（OCR）", section_path="幻灯片 19 / body-2"),
    ]
    analysis = QueryAnalysis(
        rewritten_query="基础模型页提到的平台能力包含哪些模型或数据治理能力？",
        question_type="enumeration",
        answer_focus="基础模型平台能力",
        focus_terms=["基础模型", "DeepSeek", "通义千问", "文心一言", "多模态数据治理", "文档增强解析", "知识元数据"],
    )
    draft = DraftAnswer(
        answer="基础模型页提到的平台能力包含多模态数据治理能力，以及DeepSeek、通义千问等开源大模型。",
        grounded_answer="多模态数据治理、文档增强解析、知识元数据，（deepseek、通义千问）。",
        inference_note="",
        question_type="enumeration",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    review = ReviewResult(
        passed=True,
        issues=[],
        revised_answer=draft.answer,
        revised_grounded_answer=draft.grounded_answer,
        revised_inference_note="",
        risk_level="low",
        reviewer_intervened=False,
    )

    payload = service.finalize_answer(
        "基础模型页提到的平台能力包含哪些模型或数据治理能力？",
        analysis,
        citations,
        draft,
        review,
    )

    assert payload["grounded"] is False
    assert payload["answer"] == "当前知识库中没有找到相关信息"


def test_finalize_blocks_foundation_platform_capability_answer_with_partial_output_even_if_citations_are_richer():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    citations = [
        build_hit(
            "用户空间1（知识空间、智能体空间）\n\n基础模型\n\n多模态数据治理\n\n文档增强解析\n\n知识元数据",
            section_path="幻灯片 19 / body",
        ),
        build_hit("（deepseek、通义千问、文心一言等）", section_path="幻灯片 19 / body-1"),
        build_hit(
            "完成了DeepSeek、Qwen等开源大模型的本地化部署，支持大模型基础应用学习与设备交互实践。",
            section_path="幻灯片 19 / body-3",
        ),
    ]
    analysis = QueryAnalysis(
        rewritten_query="基础模型页提到的平台能力包含哪些模型或数据治理能力？",
        question_type="enumeration",
        answer_focus="基础模型平台能力",
        focus_terms=["基础模型", "DeepSeek", "通义千问", "文心一言", "多模态数据治理", "文档增强解析", "知识元数据"],
    )
    draft = DraftAnswer(
        answer="平台能力包含多模态数据治理，并在模型方面支持DeepSeek、Qwen等开源大模型的本地化部署。",
        grounded_answer="平台应用层提供多模态数据治理能力，同时完成了DeepSeek、Qwen等开源大模型的本地化部署。",
        inference_note="从证据1提取多模态数据治理，从证据3提取DeepSeek和Qwen大模型，综合回答平台的相关能力。",
        question_type="enumeration",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    review = ReviewResult(
        passed=True,
        issues=[],
        revised_answer=draft.answer,
        revised_grounded_answer=draft.grounded_answer,
        revised_inference_note=draft.inference_note,
        risk_level="low",
        reviewer_intervened=False,
    )

    payload = service.finalize_answer(
        "基础模型页提到的平台能力包含哪些模型或数据治理能力？",
        analysis,
        citations,
        draft,
        review,
    )

    assert payload["grounded"] is False
    assert payload["answer"] == "当前知识库中没有找到相关信息。"


def test_llm_review_can_be_skipped_for_clean_or_deterministic_cases():
    service = LlmService(
        provider="openai_compatible",
        base_url="http://example.com",
        api_key="test-key",
        model="test-model",
        review_policy="auto",
        disabled=False,
    )
    citations = [build_hit("产业学院采用理事会领导下的院长负责制。", section_path="治理结构")]
    analysis = QueryAnalysis(
        rewritten_query="产业学院采用什么治理模式？",
        question_type="factoid",
        answer_focus="产业学院治理模式",
        focus_terms=["治理模式", "理事会领导下的院长负责制"],
    )
    draft = DraftAnswer(
        answer="产业学院采用理事会领导下的院长负责制。",
        grounded_answer="产业学院采用理事会领导下的院长负责制。",
        inference_note="",
        question_type="factoid",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    heuristic = service._deterministic_review("产业学院采用什么治理模式？", analysis, citations, draft)
    should_call, reason = service._should_call_llm_review("产业学院采用什么治理模式？", analysis, draft, heuristic)
    assert not should_call
    assert reason == "clean_heuristic"

    review = service.review_answer("产业学院采用什么治理模式？", analysis, citations, draft)
    assert review.raw_payload["llm_review_skipped"] is True
    assert review.raw_payload["skip_reason"] == "clean_heuristic"


def test_llm_review_kept_for_followup_focus_issue():
    service = LlmService(
        provider="openai_compatible",
        base_url="http://example.com",
        api_key="test-key",
        model="test-model",
        review_policy="auto",
        disabled=False,
    )
    citations = [build_hit("理事会作为最高决策机构。", section_path="治理结构")]
    analysis = QueryAnalysis(
        rewritten_query="产业学院的最高决策机构是什么？",
        question_type="followup",
        answer_focus="产业学院最高决策机构",
        focus_terms=["最高决策机构", "理事会"],
    )
    draft = DraftAnswer(
        answer="产业学院采用理事会领导下的院长负责制。",
        grounded_answer="理事会作为最高决策机构。",
        inference_note="",
        question_type="followup",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    heuristic = service._deterministic_review("它的最高决策机构是什么？", analysis, citations, draft)
    should_call, reason = service._should_call_llm_review("它的最高决策机构是什么？", analysis, draft, heuristic)
    assert should_call
    assert reason == "focus_alignment_issue"


def test_llm_review_skipped_for_clean_followup_answer():
    service = LlmService(
        provider="openai_compatible",
        base_url="http://example.com",
        api_key="test-key",
        model="test-model",
        review_policy="auto",
        disabled=False,
    )
    citations = [build_hit("理事会作为最高决策机构。", section_path="治理结构")]
    analysis = QueryAnalysis(
        rewritten_query="产业学院的最高决策机构是什么？",
        question_type="followup",
        answer_focus="产业学院最高决策机构",
        focus_terms=["最高决策机构", "理事会"],
    )
    draft = DraftAnswer(
        answer="最高决策机构是理事会。",
        grounded_answer="理事会作为最高决策机构。",
        inference_note="",
        question_type="followup",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    heuristic = service._deterministic_review("它的最高决策机构是什么？", analysis, citations, draft)
    should_call, reason = service._should_call_llm_review("它的最高决策机构是什么？", analysis, draft, heuristic)
    assert not should_call
    assert reason == "clean_heuristic"

    review = service.review_answer("它的最高决策机构是什么？", analysis, citations, draft)
    assert review.raw_payload["llm_review_skipped"] is True
    assert review.raw_payload["skip_reason"] == "clean_heuristic"


def test_llm_review_skipped_for_followup_verbose_but_deterministic_rewrite():
    service = LlmService(
        provider="openai_compatible",
        base_url="http://example.com",
        api_key="test-key",
        model="test-model",
        review_policy="auto",
        disabled=False,
    )
    citations = [build_hit("理事会作为最高决策机构。", section_path="治理结构")]
    analysis = QueryAnalysis(
        rewritten_query="产业学院的最高决策机构是什么？",
        question_type="followup",
        answer_focus="产业学院最高决策机构",
        focus_terms=["最高决策机构", "理事会"],
    )
    draft = DraftAnswer(
        answer="产业学院采用理事会领导下的院长负责制，其中理事会作为最高决策机构，负责重大事项决策。",
        grounded_answer="理事会作为最高决策机构。",
        inference_note="",
        question_type="followup",
        answer_focus=analysis.answer_focus,
        grounded=True,
    )
    heuristic = service._deterministic_review("它的最高决策机构是什么？", analysis, citations, draft)
    should_call, reason = service._should_call_llm_review("它的最高决策机构是什么？", analysis, draft, heuristic)
    assert not should_call
    assert reason == "followup_deterministic_rewrite_enough"


def test_generate_answer_uses_fast_path_for_pattern_and_recovered_grounded_cases():
    service = LlmService(
        provider="openai_compatible",
        base_url="http://example.com",
        api_key="test-key",
        model="test-model",
        review_policy="auto",
        disabled=False,
    )

    analysis = QueryAnalysis(
        rewritten_query="产业学院采用什么治理模式？",
        question_type="factoid",
        answer_focus="产业学院治理模式",
        focus_terms=["治理模式", "理事会领导下的院长负责制"],
    )
    draft = service.generate_answer(
        "产业学院采用什么治理模式？",
        analysis,
        [build_hit("产业学院采用理事会领导下的院长负责制。", section_path="治理结构")],
        True,
    )
    assert draft.confidence_note == "fast_path"
    assert draft.raw_payload["mode"] == "fast_path"
    assert draft.answer == "产业学院采用理事会领导下的院长负责制。"

    recovered_analysis = QueryAnalysis(
        rewritten_query="根技术体验中心的核心标语是什么？",
        question_type="factoid",
        answer_focus="根技术体验中心核心标语",
        focus_terms=["核心标语", "根生万物", "智育未来"],
    )
    recovered_draft = service.generate_answer(
        "根技术体验中心的核心标语是什么？",
        recovered_analysis,
        [build_hit("体验中心以“根生万物·智育未来”为核心标语。", section_path="文化建设")],
        False,
    )
    assert recovered_draft.confidence_note == "fast_path"
    assert recovered_draft.grounded is True
    assert recovered_draft.answer == "核心标语是“根生万物·智育未来”。"

    followup_analysis = QueryAnalysis(
        rewritten_query="协作式机械臂产品支持本地部署的大模型有哪些？",
        question_type="followup",
        answer_focus="协作式机械臂产品支持本地部署的大模型",
        focus_terms=["本地部署", "大模型", "DeepSeek", "Qwen"],
    )
    followup_draft = service.generate_answer(
        "它支持本地部署哪些大模型？",
        followup_analysis,
        [
            build_hit(
                "除以上功能外，还在运算单元中完成了DeepSeek、Qwen等开源大模型的本地化部署，"
                "支持大模型技术开发、大模型+视觉、大模型+语音、大模型+机器人等多个行业场景的应用实践。",
                section_path="产品能力 / 大模型部署",
            )
        ],
        False,
    )
    assert followup_draft.confidence_note == "fast_path"
    assert followup_draft.grounded is True
    assert followup_draft.answer == "支持本地部署DeepSeek、Qwen等开源大模型。"


def test_generate_answer_returns_insufficient_when_no_deterministic_evidence():
    service = LlmService(
        provider="openai_compatible",
        base_url="http://example.com",
        api_key="test-key",
        model="test-model",
        review_policy="auto",
        disabled=False,
    )
    ungrounded_analysis = QueryAnalysis(
        rewritten_query="华为ICT学院的学费是多少？",
        question_type="factoid",
        answer_focus="华为ICT学院学费",
        focus_terms=["学费"],
    )
    ungrounded_draft = service.generate_answer(
        "华为ICT学院的学费是多少？",
        ungrounded_analysis,
        [],
        False,
    )
    assert ungrounded_draft.confidence_note == "ungrounded_fast_path"
    assert ungrounded_draft.grounded is False
    assert "当前知识库没有直接证据" in ungrounded_draft.answer


def test_build_answer_focus_normalizes_broken_focus_fragments():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)

    focus = service._normalize_answer_focus(
        "华为根技术体验中心的核心定位主线是什么？",
        "factoid",
        "华为根技术体验中心的核心、定位主线是",
        ["华为根技术体验中心", "核心定位主线", "根技术筑基"],
    )
    assert focus == "华为根技术体验中心、核心定位主线"

    followup_focus = service._normalize_answer_focus(
        "它的最高决策机构是什么？",
        "followup",
        "产业学院最高决策机构的名称",
        ["最高决策机构", "理事会"],
    )
    assert followup_focus == "最高决策机构"

    assert (
        service._normalize_answer_focus(
            "展厅文化建设紧扣的是哪三句话？",
            "factoid",
            "展厅文化建设紧扣的是哪三、句话",
            ["根技术筑基", "产教融育人", "师范践初心"],
        )
        == "核心定位主线"
    )
    assert (
        service._normalize_answer_focus(
            "展厅重点覆盖哪四个支柱领域？",
            "enumeration",
            "展厅重点覆盖哪四个支柱领",
            ["智慧农业", "智能制造", "健康卫生", "智能教育"],
        )
        == "四个支柱方向"
    )


def test_context_block_is_compact_for_factoid_questions():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    long_text = (
        "华为人才在线官网是一站式数字化人才培养平台，集学、训、练、考、赛、证为一体，"
        "支撑学员快速获取知识和技能，同时支持讲师获取教学资源并开展教学活动。"
    )
    context = service._context_block(
        "华为人才在线官网是什么平台？",
        "factoid",
        ["华为人才在线官网", "平台"],
        [build_hit(long_text, section_path="平台介绍")],
    )
    assert "[证据 1]" in context
    assert "文件:" not in context
    assert len(context) < 160


def test_finalize_answer_strips_question_echo():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    analysis = QueryAnalysis(
        rewritten_query="广师大根技术体验中心的建设目标是什么？",
        question_type="factoid",
        answer_focus="建设目标",
        focus_terms=["建设目标", "根技术筑基", "师范践初心"],
    )
    draft = DraftAnswer(
        answer="广师大根技术体验中心的建设目标是什么？ 建设目标是根技术筑基、产教融育人、师范践初心。",
        grounded_answer="广师大根技术体验中心的建设目标是什么？ 建设目标：根技术筑基、产教融育人、师范践初心。",
        inference_note="",
        question_type="factoid",
        answer_focus="建设目标",
        grounded=True,
    )
    review = ReviewResult(
        passed=True,
        issues=[],
        revised_answer=draft.answer,
        revised_grounded_answer=draft.grounded_answer,
        revised_inference_note="",
        risk_level="low",
        reviewer_intervened=False,
    )
    payload = service.finalize_answer(
        "广师大根技术体验中心的建设目标是什么？",
        analysis,
        [build_hit("建设目标：根技术筑基、产教融育人、师范践初心。")],
        draft,
        review,
    )
    assert payload["answer"] == "建设目标是根技术筑基、产教融育人、师范践初心。"


def test_finalize_answer_normalizes_frequency_factoid():
    service = LlmService(provider="openai_compatible", base_url=None, api_key=None, model=None, disabled=True)
    analysis = QueryAnalysis(
        rewritten_query="产业学院的决策会议多久召开一次？",
        question_type="factoid",
        answer_focus="决策会议频率",
        focus_terms=["决策会议", "召开频率", "每季度"],
    )
    draft = DraftAnswer(
        answer="产业学院的决策会议每季度召开一次。 会议主要负责审定发展战略。",
        grounded_answer="资料提到理事会每季度召开一次决策会议。",
        inference_note="",
        question_type="factoid",
        answer_focus="决策会议频率",
        grounded=True,
    )
    review = ReviewResult(
        passed=True,
        issues=[],
        revised_answer=draft.answer,
        revised_grounded_answer=draft.grounded_answer,
        revised_inference_note="",
        risk_level="low",
        reviewer_intervened=False,
    )
    payload = service.finalize_answer(
        "那决策会议多久召开一次？",
        analysis,
        [build_hit("决策机制：理事会每季度召开 1 次决策会议，审定发展战略。")],
        draft,
        review,
    )
    assert payload["answer"] == "决策会议每季度召开1次。"


def test_admin_page_shows_failed_eval_examples(client, app_env: Path):
    files = [("files", ("g1.docx", build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    response = client.post(
        "/api/admin/documents",
        headers={"X-Admin-Token": "test-token"},
        files=files,
        data={"source_type": "upload", "trust_level": "internal"},
    )
    assert response.status_code == 200, response.text

    eval_dataset = app_env / "data" / "evals" / "cases.json"
    eval_dataset.parent.mkdir(parents=True, exist_ok=True)
    eval_dataset.write_text(
        json.dumps(
            [
                {
                    "id": "expected_fail_case",
                    "category": "factoid",
                    "question": "宇树 G1 知识卡说了什么？",
                    "expected_answer_keywords": ["不存在的关键词"],
                    "expected_directness": True,
                    "expected_insufficient": False,
                    "expected_grounded": True,
                }
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    eval_response = client.post("/api/admin/evals/run", headers={"X-Admin-Token": "test-token"})
    assert eval_response.status_code == 200, eval_response.text

    admin = client.get("/admin?token=test-token")
    assert admin.status_code == 200, admin.text
    html = admin.text
    assert "快速回归" in html
    assert "全量验收" in html
    assert "失败样例直达" in html
    assert "expected_fail_case#1" in html
    assert "查看本次链路" in html


def test_ragflow_sync_reuses_current_version(client):
    class FakeEmbeddingService:
        def embed_documents(self, texts):
            return [[0.1, 0.2, 0.3] for _ in texts]

    class FakeRagflowClient:
        def list_dataset_documents(self, dataset_id, page_size=100):
            return [{"id": "doc-1", "name": "知识库材料.pdf"}]

        def list_document_chunks(self, dataset_id, document_id, page_size=100):
            return [
                {"id": "chunk-1", "content": "第一段知识点", "positions": [[1, 0, 0, 0, 0]]},
                {"id": "chunk-2", "content": "第二段知识点", "positions": [[2, 0, 0, 0, 0]]},
            ]

    from app.services.ragflow_sync import RagflowSyncService

    container = client.app.state.container
    sync_service = RagflowSyncService(
        repository=container.repository,
        ragflow_client=FakeRagflowClient(),
        embedding_service=FakeEmbeddingService(),
        vector_store=container.vector_store,
        upload_dir=container.settings.upload_dir,
    )

    first = sync_service.sync_dataset("ds-1")
    second = sync_service.sync_dataset("ds-1")

    assert first["document_count"] == 1
    assert first["chunk_count"] == 2
    assert second["document_count"] == 1
    assert second["chunk_count"] == 2

    docs = container.repository.list_documents()
    target = next(doc for doc in docs if doc["title"] == "知识库材料")
    assert target["chunk_count"] == 2

    with container.db.connect() as conn:
        rows = conn.execute(
            """
            SELECT version_number, stored_path, parser_name, chunk_count
            FROM document_versions
            WHERE document_id = ?
            ORDER BY version_number ASC
            """,
            (target["document_id"],),
        ).fetchall()
    assert len(rows) == 1
    assert rows[0]["version_number"] == 1
    assert rows[0]["parser_name"] == "ragflow_import"
    assert rows[0]["chunk_count"] == 2


def test_ragflow_sync_reuses_existing_non_current_version(client):
    class FakeEmbeddingService:
        def embed_documents(self, texts):
            return [[0.1, 0.2, 0.3] for _ in texts]

    class FakeRagflowClient:
        def list_dataset_documents(self, dataset_id, page_size=100):
            return [{"id": "doc-1", "name": "知识库材料.pdf"}]

        def list_document_chunks(self, dataset_id, document_id, page_size=100):
            return [{"id": "chunk-1", "content": "第一段知识点", "positions": [[1, 0, 0, 0, 0]]}]

    from app.services.ragflow_sync import RagflowSyncService

    container = client.app.state.container
    sync_service = RagflowSyncService(
        repository=container.repository,
        ragflow_client=FakeRagflowClient(),
        embedding_service=FakeEmbeddingService(),
        vector_store=container.vector_store,
        upload_dir=container.settings.upload_dir,
    )

    sync_service.sync_dataset("ds-1")
    docs = container.repository.list_documents()
    target = next(doc for doc in docs if doc["title"] == "知识库材料")

    replacement = container.repository.create_document_version(
        document_id=target["document_id"],
        original_filename="知识库材料.pdf",
        stored_path="data/uploads/mock/replacement.pdf",
        file_size=1,
        sha256="replacement-sha",
    )
    container.repository.update_version_status(
        replacement["id"],
        parser_status="completed",
        index_status="completed",
        parser_name="manual_override",
        ocr_used=False,
        warning_text="",
        chunk_count=0,
    )

    result = sync_service.sync_dataset("ds-1")
    assert result["document_count"] == 1

    refreshed = next(doc for doc in container.repository.list_documents() if doc["document_id"] == target["document_id"])
    with container.db.connect() as conn:
        rows = conn.execute(
            """
            SELECT version_number, stored_path
            FROM document_versions
            WHERE document_id = ?
            ORDER BY version_number ASC
            """,
            (target["document_id"],),
        ).fetchall()
    assert len(rows) == 2
    assert refreshed["current_version"] != replacement["id"]
    assert refreshed["chunk_count"] == 1


def test_reranker_runtime_error_falls_back_to_lexical():
    reranker = RerankerService("dummy-model", use_stub=False)

    class BrokenBackend:
        def compute_score(self, pairs, batch_size=1):
            raise NotImplementedError("Cannot copy out of meta tensor")

    reranker._backend = BrokenBackend()
    hits = [
        build_hit("功能全面、性能优异、全球共享。", section_path="优势"),
        build_hit("无关内容。", section_path="其他"),
    ]

    result = reranker.rerank("华为人才的优势是什么", hits)
    assert result
    assert result[0].rerank_score >= result[-1].rerank_score
    assert reranker._backend == "stub"


def test_cleanup_ragflow_versions_api_removes_duplicate_versions(client):
    class FakeEmbeddingService:
        def embed_documents(self, texts):
            return [[0.1, 0.2, 0.3] for _ in texts]

    class FakeRagflowClient:
        def list_dataset_documents(self, dataset_id, page_size=100):
            return [{"id": "doc-1", "name": "知识库材料.pdf"}]

        def list_document_chunks(self, dataset_id, document_id, page_size=100):
            return [{"id": "chunk-1", "content": "第一段知识点", "positions": [[1, 0, 0, 0, 0]]}]

    from app.services.ragflow_sync import RagflowSyncService

    container = client.app.state.container
    sync_service = RagflowSyncService(
        repository=container.repository,
        ragflow_client=FakeRagflowClient(),
        embedding_service=FakeEmbeddingService(),
        vector_store=container.vector_store,
        upload_dir=container.settings.upload_dir,
    )

    sync_service.sync_dataset("ds-1")
    docs = container.repository.list_documents()
    target = next(doc for doc in docs if doc["title"] == "知识库材料")
    duplicate = container.repository.create_document_version(
        document_id=target["document_id"],
        original_filename="知识库材料.pdf",
        stored_path="ragflow://ds-1/doc-1/知识库材料.pdf",
        file_size=0,
        sha256="same-ragflow-sha",
    )
    container.repository.update_version_status(
        duplicate["id"],
        parser_status="completed",
        index_status="completed",
        parser_name="ragflow_import",
        ocr_used=False,
        warning_text="",
        chunk_count=1,
    )

    response = client.post("/api/admin/ragflow/cleanup", headers={"X-Admin-Token": "test-token"})
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["removed_versions"] == 1
    assert payload["affected_documents"] == 1

    with container.db.connect() as conn:
        rows = conn.execute(
            "SELECT version_number FROM document_versions WHERE document_id = ? ORDER BY version_number ASC",
            (target["document_id"],),
        ).fetchall()
    assert [row["version_number"] for row in rows] == [2]


def test_evaluation_service_can_run_via_http_api(client, monkeypatch):
    eval_dataset = Path(client.app.state.container.settings.eval_dataset_path)
    eval_dataset.parent.mkdir(parents=True, exist_ok=True)
    eval_dataset.write_text(
        json.dumps(
            [
                {
                    "id": "api_eval_case",
                    "category": "factoid",
                    "question": "华为人才在线官网是什么平台？",
                    "expected_answer_keywords": ["一站式数字化人才培养平台"],
                    "expected_directness": True,
                    "expected_insufficient": False,
                    "expected_grounded": True,
                }
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    monkeypatch.setattr(client.app.state.container.settings, "eval_api_base_url", "http://testserver")
    client.app.state.container.evaluation_service.http_client = client
    payload = client.app.state.container.evaluation_service.run(dataset_path=eval_dataset)
    assert payload["summary"]["total_turns"] == 1
    assert "report_path" in payload
    assert "formal_summary_path" in payload
    assert Path(payload["formal_summary_path"]).exists()


def test_evaluation_service_builds_route_conflict_formal_block(app_env: Path):
    settings = Settings(admin_token="test-token")
    service = EvaluationService(settings, repository=None, chat_service=None)  # type: ignore[arg-type]
    report = service._evaluate_turn(
        case_id="ppt-company-p0-05",
        category="new_ppt_enumeration",
        turn_index=1,
        question="基础环境页提到了哪些基础设施或平台能力？请列举至少4项。",
        answer_payload=type(
            "Payload",
            (),
            {
                "answer": "包括旧资料里的平台和硬件项",
                "grounded_answer": "",
                "inference_note": "",
                "grounded": True,
                "citations": [
                    RetrievalHit(
                        chunk_id="chunk-1",
                        document_id="doc-1",
                        version_id="ver-1",
                        file_name="华为ICT学院手册 2024-2025.pdf",
                        page_or_slide="46",
                        section_path="5 > 46 > 546 > 352 > 499",
                        snippet="旧资料片段",
                        markdown_text="旧资料片段",
                        plain_text="旧资料片段",
                        trust_level="internal",
                        source_type="api_eval",
                        fusion_score=1.0,
                        rerank_score=1.0,
                    )
                ],
                "reviewer_intervened": True,
                "fallback_used": True,
                "latency_ms": 10,
                "conversation_id": "conv-1",
                "answer_run_id": "run-1",
                "question_type": "enumeration",
                "answer_focus": "基础环境、基础设施",
                "review_issues": [],
            },
        )(),
        expected_files=["【公司介绍】轩辕网络公司介绍202606.pptx"],
        expected_section_keywords=["基础环境"],
        expected_answer_keywords=[],
        forbidden_answer_keywords=[],
        max_answer_length=130,
        expected_question_type="enumeration",
        expected_directness=True,
        expected_insufficient=True,
        expected_grounded=False,
    )
    formal_summary = EvaluationService.build_formal_summary(
        dataset_cases=[
            {
                "id": "ppt-company-p0-05",
                "expected_result_mode": "must_block",
                "blocking_is_correct_if_any": "route_conflict",
            }
        ],
        reports=[report],
        source_report="data/evals/results/eval_fake.json",
    )

    assert formal_summary["formal_summary"]["correct_block"] == 1
    assert formal_summary["formal_summary"]["wrong_release"] == 0
    assert formal_summary["formal_reports"][0]["formal_bucket"] == "correct_block"
    assert formal_summary["formal_reports"][0]["citation_match"] is False


def test_evaluation_service_builds_failed_case_dataset(app_env: Path):
    source_dataset = app_env / "data" / "evals" / "source_cases.json"
    source_dataset.parent.mkdir(parents=True, exist_ok=True)
    source_dataset.write_text(
        json.dumps(
            [
                {
                    "id": "case_ok",
                    "category": "factoid",
                    "question": "问题一",
                },
                {
                    "id": "case_fail",
                    "category": "multi_turn",
                    "turns": [{"question": "问题二"}, {"question": "追问二"}],
                },
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    latest_result = {
        "dataset": str(source_dataset),
        "reports": [
            {"case_id": "case_ok", "passed": True},
            {"case_id": "case_fail", "passed": False},
            {"case_id": "case_fail", "passed": True},
        ],
    }
    output_path = app_env / "data" / "evals" / "results" / "failed_only.json"
    generated = EvaluationService.build_failed_case_dataset(latest_result, output_path=output_path)
    assert generated["case_count"] == 1
    assert generated["failed_case_ids"] == ["case_fail"]
    payload = json.loads(output_path.read_text(encoding="utf-8"))
    assert len(payload) == 1
    assert payload[0]["id"] == "case_fail"
    assert len(payload[0]["turns"]) == 2


def test_formal_summary_marks_blocked_answer_as_correct_block_even_if_report_failed():
    report = EvaluationService(
        Settings(),
        repository=None,  # type: ignore[arg-type]
        chat_service=None,  # type: ignore[arg-type]
    )._evaluate_turn(
        case_id="ppt-company-p0-03",
        category="new_ppt_summary",
        turn_index=1,
        question="轩辕网络的业务架构在 PPT 中是如何概括的？",
        answer_payload=type(
            "Payload",
            (),
            {
                "answer": "当前知识库中没有找到相关操作步骤",
                "grounded_answer": "",
                "inference_note": "",
                "grounded": False,
                "citations": [
                    RetrievalHit(
                        chunk_id="chunk-1",
                        document_id="doc-1",
                        version_id="ver-1",
                        file_name="【公司介绍】轩辕网络公司介绍202606.pptx",
                        page_or_slide="11",
                        section_path="幻灯片 11 / 幻灯片 11 / body",
                        snippet="业务架构：有产懂教，双轮驱动",
                        markdown_text="业务架构：有产懂教，双轮驱动",
                        plain_text="业务架构：有产懂教，双轮驱动",
                        trust_level="internal",
                        source_type="api_eval",
                        fusion_score=1.0,
                        rerank_score=1.0,
                    )
                ],
                "reviewer_intervened": True,
                "fallback_used": False,
                "latency_ms": 10,
                "conversation_id": "conv-1",
                "answer_run_id": "run-1",
                "question_type": "procedure",
                "answer_focus": "业务架构",
                "review_issues": [],
            },
        )(),
        expected_files=["【公司介绍】轩辕网络公司介绍202606.pptx"],
        expected_section_keywords=["业务架构", "双轮驱动"],
        expected_answer_keywords=[],
        forbidden_answer_keywords=[],
        max_answer_length=90,
        expected_question_type="factoid",
        expected_directness=True,
        expected_insufficient=True,
        expected_grounded=False,
    )

    formal_summary = EvaluationService.build_formal_summary(
        dataset_cases=[
            {
                "id": "ppt-company-p0-03",
                "expected_result_mode": "must_block",
                "blocking_is_correct_if_any": "grounding_insufficient",
            }
        ],
        reports=[report],
        source_report="data/evals/results/eval_fake.json",
    )

    assert report.passed is False
    assert report.insufficient_match is True
    assert report.grounded_match is True
    assert formal_summary["formal_summary"]["correct_block"] == 1
    assert formal_summary["formal_summary"]["wrong_release"] == 0
    assert formal_summary["formal_reports"][0]["formal_bucket"] == "correct_block"


def test_formal_summary_marks_compact_answer_pass_without_section_match():
    report = EvaluationService(
        Settings(),
        repository=None,  # type: ignore[arg-type]
        chat_service=None,  # type: ignore[arg-type]
    )._evaluate_turn(
        case_id="ppt-company-p0-02",
        category="new_ppt_factoid",
        turn_index=1,
        question="公司概况页提到轩辕网络深耕教育多少年、专注什么方向？",
        answer_payload=type(
            "Payload",
            (),
            {
                "answer": "轩辕网络深耕教育28年，专注产教融合方向。",
                "grounded_answer": "",
                "inference_note": "",
                "grounded": True,
                "citations": [
                    RetrievalHit(
                        chunk_id="chunk-1",
                        document_id="doc-1",
                        version_id="ver-1",
                        file_name="【公司介绍】轩辕网络公司介绍202606.pptx",
                        page_or_slide="3",
                        section_path="幻灯片 3 / 幻灯片 3 / body",
                        snippet="28年教育深耕者 专注产教融合",
                        markdown_text="28年教育深耕者 专注产教融合",
                        plain_text="28年教育深耕者 专注产教融合",
                        trust_level="internal",
                        source_type="api_eval",
                        fusion_score=1.0,
                        rerank_score=1.0,
                    )
                ],
                "reviewer_intervened": False,
                "fallback_used": False,
                "latency_ms": 10,
                "conversation_id": "conv-1",
                "answer_run_id": "run-1",
                "question_type": "factoid",
                "answer_focus": "公司概况",
                "review_issues": [],
            },
        )(),
        expected_files=["【公司介绍】轩辕网络公司介绍202606.pptx"],
        expected_section_keywords=["公司概况", "28年教育深耕者", "专注产教融合"],
        expected_answer_keywords=["28年", "产教融合"],
        forbidden_answer_keywords=["华为ICT学院"],
        max_answer_length=50,
        expected_question_type="factoid",
        expected_directness=True,
        expected_insufficient=None,
        expected_grounded=True,
    )

    formal_summary = EvaluationService.build_formal_summary(
        dataset_cases=[
            {
                "id": "ppt-company-p0-02",
                "expected_result_mode": "must_answer_compact",
                "blocking_is_correct_if_any": "none",
            }
        ],
        reports=[report],
        source_report="data/evals/results/eval_fake.json",
    )

    assert report.passed is False
    assert report.section_match is False
    assert report.citation_match is True
    assert report.keyword_match is True
    assert formal_summary["formal_summary"]["answer_pass"] == 1
    assert formal_summary["formal_summary"]["wrong_block"] == 0
    assert formal_summary["formal_reports"][0]["formal_bucket"] == "answer_pass"


def test_admin_can_run_failed_eval_job(client, app_env: Path):
    client.app.state.container.settings.eval_api_base_url = None
    eval_dataset = app_env / "data" / "evals" / "cases.json"
    eval_dataset.parent.mkdir(parents=True, exist_ok=True)
    eval_dataset.write_text(
        json.dumps(
            [
                {
                    "id": "expected_fail_case",
                    "category": "factoid",
                    "question": "宇树 G1 知识卡说了什么？",
                    "expected_answer_keywords": ["不存在的关键词"],
                    "expected_directness": True,
                    "expected_insufficient": False,
                    "expected_grounded": True,
                }
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    response = client.post("/api/admin/evals/run", headers={"X-Admin-Token": "test-token"})
    assert response.status_code == 200, response.text

    rerun = client.post("/api/admin/evals/run-failures", headers={"X-Admin-Token": "test-token"})
    assert rerun.status_code == 200, rerun.text
    payload = rerun.json()
    assert payload["status"] == "queued"
    assert payload["case_count"] == 1
    assert payload["dataset"].endswith("latest_failed_cases.json")
